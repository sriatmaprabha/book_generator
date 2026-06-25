"""
Workflow orchestrator — runs the 6-agent pipeline with session state tracking.

Pipeline:
  1. Intake        (sequential)   — validate config
  2. Architect     (sequential)   — design blueprint
  3. Researcher    (parallel)     — gather material per chapter
  3.5. YouTube     (sequential)   — match @SriNithyananda videos per chapter
  4. Writer        (batched)      — write chapters with rolling context
  5. Editor        (parallel)     — polish all chapters
  5.5. QA Review  (sequential)   — score chapters, trigger rewrites if needed
  5.55. Q&A Gen   (parallel)     — generate Q&A pairs per chapter
  5.6. Frontmatter (parallel)    — generate foreword + benediction
  6. Designer      (sequential)   — produce .docx

Every agent call is traced via Tracer — full input/output logged to
  output/<run>/agent_traces.log   (human-readable)
  output/<run>/agent_traces.jsonl (machine-readable)
"""

from __future__ import annotations

import asyncio
import json
import re
from datetime import datetime
from pathlib import Path
from textwrap import dedent
from typing import Any, Dict, List, Optional

from models import (
    AdminReview,
    BookBlueprint,
    BookConfig,
    BookSectionGroup,
    ChapterBrief,
    ChapterDraft,
    ChapterQA,
    CompiledBookMetadata,
    EditedChapter,
    QAPair,
    ResearchPacket,
    SourceLink,
)
from config import AppConfig, load_config
from models import BookQAReview, ChapterQAResult, ContentDraft, StoryDraft
from agents import (
    architect_agent,
    benediction_agent,
    combiner_writer_agent,
    content_writer_agent,
    designer_agent,
    editor_agent,
    foreword_agent,
    intake_agent,
    qa_agent,
    qa_generator_agent,
    targeted_qa_agent,
    researcher_agent,
    story_writer_agent,
    writer_agent,
    STORY_FORMAT_CYCLE,
)
from tracing import Tracer
import youtube_matcher as yt_matcher

from agno.tools.mcp import MCPTools


# ── MCP Connection ──────────────────────────────────────────────────────────

class McpConnection:
    """
    Manages connections to MCP servers defined in config.yaml.

    Creates per-agent connection pools to avoid contention when
    multiple agents run in parallel.
    """

    def __init__(self):
        self._shared: List[MCPTools] = []       # for sequential agents (Architect)
        self._pools: List[List[MCPTools]] = []   # per-agent pools for parallel use
        self._all: List[MCPTools] = []           # everything, for cleanup
        self._sources: list = []
        self._connected = False

    async def _connect_one_set(self, app_config: AppConfig, label: str = "") -> List[MCPTools]:
        """Connect to all configured MCP sources once. Returns list of MCPTools."""
        connections = []
        for source in app_config.mcp_sources:
            try:
                mcp = MCPTools(
                    transport=source.transport,
                    url=source.url,
                    timeout_seconds=120,
                )
                await mcp.connect()
                connections.append(mcp)
                self._all.append(mcp)
            except Exception as exc:
                print(f"  WARNING: Failed to connect to {source.name}{label}: {exc}")
        return connections

    async def connect(self, app_config: AppConfig) -> List[MCPTools]:
        """Connect shared set for sequential agents."""
        if self._connected:
            return self._shared
        self._sources = app_config.mcp_sources
        for source in app_config.mcp_sources:
            print(f"  Connecting to MCP: {source.name} ({source.url})...")
            try:
                mcp = MCPTools(
                    transport=source.transport,
                    url=source.url,
                    timeout_seconds=120,
                )
                await mcp.connect()
                self._shared.append(mcp)
                self._all.append(mcp)
                print(f"  Connected: {source.name}")
            except Exception as exc:
                print(f"  WARNING: Failed to connect to {source.name}: {exc}")
        self._connected = True
        return self._shared

    async def get_pool_for_agent(self, app_config: AppConfig) -> List[MCPTools]:
        """
        Create a FRESH set of MCP connections for one parallel agent.
        Each parallel agent gets its own connections to avoid contention.
        """
        pool = await self._connect_one_set(app_config, label=" (parallel pool)")
        self._pools.append(pool)
        return pool

    async def close(self):
        """Close ALL MCP connections (shared + pools)."""
        for mcp in self._all:
            try:
                await mcp.close()
            except Exception:
                pass
        self._shared.clear()
        self._pools.clear()
        self._all.clear()
        self._connected = False

    @property
    def tools(self) -> List[MCPTools]:
        return self._shared


# ── Utilities ───────────────────────────────────────────────────────────────

def safe_print(text: str) -> None:
    """Print with fallback for Windows cp1252 console encoding issues."""
    try:
        print(text)
    except UnicodeEncodeError:
        print(text.encode("ascii", "replace").decode())


def slugify(value: str) -> str:
    cleaned = "".join(ch if ch.isalnum() or ch in (" ", "-", "_") else "_" for ch in value)
    return "_".join(cleaned.split())[:80] or "book"


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def write_json(path: Path, data: Any) -> None:
    ensure_dir(path.parent)
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


def write_text(path: Path, content: str) -> None:
    ensure_dir(path.parent)
    path.write_text(content, encoding="utf-8")


def count_words(text: str) -> int:
    return len(text.split())


# ── Session State ──────────────────────────────────────────────────────────

class PipelineState:
    """Central state object passed through the pipeline — the handoff ledger."""

    def __init__(self, book_config: BookConfig, output_dir: Path, tracer: Tracer):
        self.config = book_config
        self.output_dir = output_dir
        self.tracer = tracer
        self.blueprint: Optional[BookBlueprint] = None
        self.research: Dict[int, ResearchPacket] = {}
        self.drafts: Dict[int, ChapterDraft] = {}
        self.rolling_summaries: Dict[int, str] = {}
        self.edited: Dict[int, EditedChapter] = {}
        self.metadata: Optional[CompiledBookMetadata] = None
        self.docx_path: Optional[Path] = None
        # story_formats[chapter_number] = "A" | "B" | "C" | "D"
        self.story_formats: Dict[int, str] = {}
        # qa_locked: chapter numbers that passed QA and must not be re-evaluated
        self.qa_locked: set = set()
        # chapter_qa: Q&A pairs per chapter
        self.chapter_qa: Dict[int, ChapterQA] = {}
        # all_source_links: YouTube links collected from research packets
        self.all_source_links: List[SourceLink] = []

        # Progress tracking
        self.status = {
            "phase": "init",
            "chapters_researched": [],
            "chapters_drafted": [],
            "chapters_edited": [],
            "current_step": "",
            "errors": [],
        }

    def to_dict(self) -> dict:
        return {
            "config": self.config.model_dump(),
            "blueprint": self.blueprint.model_dump() if self.blueprint else None,
            "research": {k: v.model_dump() for k, v in self.research.items()},
            "drafts": {k: v.model_dump() for k, v in self.drafts.items()},
            "rolling_summaries": self.rolling_summaries,
            "edited": {k: v.model_dump() for k, v in self.edited.items()},
            "metadata": self.metadata.model_dump() if self.metadata else None,
            "story_formats": self.story_formats,
            "qa_locked": list(self.qa_locked),
            "status": self.status,
        }

    def save_snapshot(self, label: str = "state") -> Path:
        path = self.output_dir / f"{label}.json"
        write_json(path, self.to_dict())
        return path


# ── Phase 1: Intake ────────────────────────────────────────────────────────

async def run_intake(
    raw_input: dict,
    app_config: AppConfig,
    tracer: Tracer,
) -> BookConfig:
    """Validate and enrich user input into a BookConfig."""
    print("\n  Phase 1: Intake")
    print("  " + "-" * 50)

    # Try direct Pydantic validation first — fast path (no LLM needed)
    try:
        book_config = BookConfig.model_validate(raw_input)
        tracer._console_log("Intake", "Validator", "SKIP", "Direct Pydantic validation OK — no LLM call needed")
        return book_config
    except Exception as validation_err:
        tracer._console_log("Intake", "Validator", "FALLBACK", f"Validation failed: {validation_err}")

    # Fall back to the Intake Agent for fuzzy / incomplete input
    agent = intake_agent(app_config)
    prompt = dedent(f"""\
        Validate and complete the following book configuration.
        Fill in any missing fields with sensible defaults.
        Return a complete BookConfig JSON.

        Raw input:
        {json.dumps(raw_input, indent=2)}
    """)
    response = await tracer.traced_arun(agent, prompt, phase="Intake", output_schema=BookConfig)
    if isinstance(response.content, BookConfig):
        return response.content

    # Last resort: try to parse the text response
    text = str(response.content or "")
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        book_config = BookConfig.model_validate_json(match.group())
        return book_config

    raise RuntimeError("Intake agent could not produce a valid BookConfig.")


# ── Phase 2: Architect ─────────────────────────────────────────────────────

async def run_architect(
    state: PipelineState,
    tools: Optional[List[Any]] = None,
    app_config: Optional[AppConfig] = None,
) -> BookBlueprint:
    """Design the book blueprint."""
    cfg = app_config or load_config()
    tracer = state.tracer
    print("\n  Phase 2: Architect")
    print("  " + "-" * 50)
    state.status["phase"] = "architect"
    state.status["current_step"] = "designing_blueprint"

    agent = architect_agent(state.config, tools=tools, config=cfg)
    prompt = dedent(f"""\
        Design a complete book blueprint for:
        Title: {state.config.title}
        Chapters: {state.config.num_chapters}

        Output a BookBlueprint with all chapter briefs.
    """)

    response = await tracer.traced_arun(agent, prompt, phase="Architect", output_schema=BookBlueprint)
    if isinstance(response.content, BookBlueprint):
        blueprint = response.content
    else:
        text = str(response.content or "")
        match = re.search(r"\{.*\}", text, re.DOTALL)
        if match:
            blueprint = BookBlueprint.model_validate_json(match.group())
        else:
            raise RuntimeError("Architect agent did not produce a valid BookBlueprint.")

    # Validate and report section structure
    type_counts = {}
    for ch in blueprint.chapters:
        type_counts[ch.chapter_type] = type_counts.get(ch.chapter_type, 0) + 1

    # If the architect didn't assign types (older runs / fallback), auto-assign sensibly
    if not any(ch.chapter_type != "main" for ch in blueprint.chapters):
        total = len(blueprint.chapters)
        for i, ch in enumerate(blueprint.chapters):
            if i == 0:
                ch.chapter_type = "preface"
                ch.target_word_count = cfg.workflow.preface_words
            elif i == 1:
                ch.chapter_type = "introduction"
            elif i == total - 1:
                ch.chapter_type = "conclusion"
                ch.target_word_count = cfg.workflow.conclusion_words
            elif i <= 2:
                ch.chapter_type = "preliminary"
                ch.section_group = "Part I: The Ground of Being"
            else:
                ch.chapter_type = "main"
                ch.section_group = "Part II: The Core Teaching"
        print("  Note: chapter types auto-assigned (Architect did not provide them)")

    section_summary = " | ".join(f"{t}:{c}" for t, c in sorted(type_counts.items()))
    print(f"  Chapter types: {section_summary}")
    if blueprint.sections:
        for sec in blueprint.sections:
            print(f"    {sec.label}: chapters {sec.chapter_numbers}")

    state.blueprint = blueprint
    write_json(state.output_dir / "blueprint.json", blueprint.model_dump())
    write_text(state.output_dir / "blueprint.md", _render_blueprint_markdown(blueprint))
    print(f"  Blueprint created: {len(blueprint.chapters)} chapters")
    state.save_snapshot()
    return blueprint


def _render_blueprint_markdown(bp: BookBlueprint) -> str:
    lines = [f"# {bp.book_title}", "", f"**Thematic Arc:** {bp.thematic_arc}", ""]
    if bp.recurring_motifs:
        lines.append(f"**Motifs:** {', '.join(bp.recurring_motifs)}")
        lines.append("")
    if bp.voice_notes:
        lines.append(f"**Voice Notes:** {bp.voice_notes}")
        lines.append("")
    lines.append("---\n")
    for ch in bp.chapters:
        lines.extend([
            f"## Chapter {ch.chapter_number}: {ch.title}",
            f"**Synopsis:** {ch.synopsis}",
            f"**Story Seed:** {ch.story_seed}",
            f"**Narrative Arc:** {ch.narrative_arc}",
            f"**Teaching Points:** {', '.join(ch.teaching_points)}",
            f"**Verse References:** {', '.join(ch.verse_references) if ch.verse_references else 'TBD'}",
            f"**Humor Seed:** {ch.humor_seed}",
            f"**Bridge to Next:** {ch.bridge_to_next}",
            "",  # target word count omitted from output
            "",
        ])
    return "\n".join(lines)


# ── Phase 3: Research (parallel) ───────────────────────────────────────────

async def run_research(
    state: PipelineState,
    tools: Optional[List[Any]] = None,
    app_config: Optional[AppConfig] = None,
) -> Dict[int, ResearchPacket]:
    """Research all chapters — runs in parallel if configured."""
    cfg = app_config or load_config()
    tracer = state.tracer
    print("\n  Phase 3: Research")
    print("  " + "-" * 50)
    state.status["phase"] = "research"

    assert state.blueprint is not None, "Blueprint must be set before research."

    # Get the MCP connection manager from pipeline context
    mcp_conn: Optional[McpConnection] = getattr(state, '_mcp_conn', None)

    async def research_one(brief: ChapterBrief, agent_tools: Optional[List[Any]] = None) -> ResearchPacket:
        state.status["current_step"] = f"researching_ch_{brief.chapter_number}"
        agent = researcher_agent(state.config, brief, tools=agent_tools, config=cfg)
        prompt = f"Research material for chapter {brief.chapter_number}: {brief.title}"
        response = await tracer.traced_arun(
            agent, prompt,
            phase=f"Research-Ch{brief.chapter_number}",
            output_schema=ResearchPacket,
        )
        if isinstance(response.content, ResearchPacket):
            packet = response.content
        else:
            packet = ResearchPacket(chapter_number=brief.chapter_number)
        packet.chapter_number = brief.chapter_number
        state.status["chapters_researched"].append(brief.chapter_number)
        return packet

    if cfg.workflow.parallel_research:
        # Each parallel researcher gets its OWN MCP connections to avoid contention
        async def research_with_own_pool(brief: ChapterBrief) -> ResearchPacket:
            if mcp_conn and cfg:
                pool = await mcp_conn.get_pool_for_agent(cfg)
                return await research_one(brief, agent_tools=pool)
            return await research_one(brief, agent_tools=tools)

        tasks = [research_with_own_pool(brief) for brief in state.blueprint.chapters]
        packets = await asyncio.gather(*tasks, return_exceptions=True)
    else:
        packets = []
        for brief in state.blueprint.chapters:
            packets.append(await research_one(brief, agent_tools=tools))

    for packet in packets:
        if isinstance(packet, Exception):
            state.status["errors"].append(str(packet))
            continue
        state.research[packet.chapter_number] = packet
        # Collect YouTube source links from each research packet
        state.all_source_links.extend(packet.source_links)

    research_dir = ensure_dir(state.output_dir / "research")
    for num, packet in state.research.items():
        write_json(research_dir / f"ch_{num:02d}_research.json", packet.model_dump())

    # Export all source links as a standalone file for QR code generation
    if state.all_source_links:
        write_json(
            state.output_dir / "source_links.json",
            [lnk.model_dump() for lnk in state.all_source_links],
        )
        print(f"  Source links extracted: {len(state.all_source_links)} YouTube URLs")

    print(f"  Research complete: {len(state.research)} chapters")
    state.save_snapshot()
    return state.research


# ── Phase 3.5: YouTube Video Matching ─────────────────────────────────────

async def run_youtube_matching(
    state: PipelineState,
    app_config: Optional[AppConfig] = None,
) -> None:
    """
    Match YouTube videos from @SriNithyananda for each chapter.

    Uses the YouTube Data API v3 to search for videos whose title/description
    matches the chapter's teaching_points.  Results are appended to
    state.all_source_links and re-exported as source_links.json.

    Skipped automatically when youtube_enabled = false or the API key is absent.
    """
    cfg = app_config or load_config()
    wf = cfg.workflow

    if not wf.youtube_enabled:
        print("\n  Phase 3.5: YouTube Matching — SKIPPED (youtube_enabled=false)")
        return

    api_key = wf.youtube_api_key
    if not api_key:
        print("\n  Phase 3.5: YouTube Matching — SKIPPED (YOUTUBE_API_KEY not set)")
        return

    print("\n  Phase 3.5: YouTube Video Matching")
    print("  " + "-" * 50)
    print(f"  Channel: {wf.youtube_channel_url}")
    state.status["phase"] = "youtube_matching"

    assert state.blueprint is not None, "Blueprint must be set before YouTube matching."

    # Build the per-chapter data list expected by youtube_matcher
    chapters_data = [
        {
            "chapter_number": brief.chapter_number,
            "title": brief.title,
            "teaching_points": brief.teaching_points,
        }
        for brief in state.blueprint.chapters
    ]

    try:
        new_links = yt_matcher.match_videos_for_chapters(
            chapters_data=chapters_data,
            api_key=api_key,
            pause_between=0.4,
        )
    except Exception as exc:
        print(f"  WARNING: YouTube matching failed — {exc}")
        print("  Continuing without YouTube source links.")
        return

    if not new_links:
        print("  YouTube: no matching videos found.")
        return

    # De-duplicate against links already collected from research packets
    existing_urls = {lnk.url for lnk in state.all_source_links}
    fresh = [lnk for lnk in new_links if lnk.url not in existing_urls]
    state.all_source_links.extend(fresh)

    # Export the updated full list
    write_json(
        state.output_dir / "source_links.json",
        [lnk.model_dump() for lnk in state.all_source_links],
    )

    print(f"  YouTube matching complete: {len(fresh)} new link(s) added "
          f"({len(state.all_source_links)} total)")
    state.save_snapshot()


# ── Phase 4: Writing (batched sequential) ──────────────────────────────────

async def run_writing(
    state: PipelineState,
    app_config: Optional[AppConfig] = None,
) -> Dict[int, ChapterDraft]:
    """Write all chapters in batches, maintaining rolling summaries."""
    cfg = app_config or load_config()
    tracer = state.tracer
    print("\n  Phase 4: Writing")
    print("  " + "-" * 50)
    state.status["phase"] = "writing"

    assert state.blueprint is not None, "Blueprint must be set before writing."
    batch_size = cfg.workflow.writer_batch_size
    chapters = state.blueprint.chapters
    draft_dir = ensure_dir(state.output_dir / "drafts")
    parts_dir = ensure_dir(state.output_dir / "drafts" / "parts")

    # Pre-assign story formats deterministically — A, B, C, D, A, B, C, D...
    # This guarantees variety regardless of what the LLM would choose on its own.
    for idx, brief in enumerate(chapters):
        state.story_formats[brief.chapter_number] = STORY_FORMAT_CYCLE[idx % len(STORY_FORMAT_CYCLE)]

    format_preview = "  ".join(
        f"Ch{n}={f}" for n, f in sorted(state.story_formats.items())
    )
    print(f"\n    Story format schedule: {format_preview}")

    prior_summaries: List[str] = []

    for i in range(0, len(chapters), batch_size):
        batch = chapters[i : i + batch_size]
        print(f"\n    Writing batch {i // batch_size + 1}: chapters {batch[0].chapter_number}-{batch[-1].chapter_number}")

        for brief in batch:
            state.status["current_step"] = f"writing_ch_{brief.chapter_number}"
            research = state.research.get(brief.chapter_number)
            priors = prior_summaries if brief.chapter_number > 1 else None

            # Collect formats already used so the agent has full context
            assigned_fmt = state.story_formats[brief.chapter_number]
            formats_so_far = [
                state.story_formats[n]
                for n in sorted(state.story_formats)
                if n < brief.chapter_number
            ]

            # ── Step 1: Story Writer + Content Writer in PARALLEL ──────
            print(f"      Ch {brief.chapter_number}: Story [{assigned_fmt}] + Content writers (parallel)...")

            story_agent = story_writer_agent(
                state.config, state.blueprint, brief,
                research=research, prior_summaries=priors,
                assigned_format=assigned_fmt,
                formats_used=formats_so_far,
                config=cfg,
            )
            content_agent = content_writer_agent(
                state.config, state.blueprint, brief,
                research=research, prior_summaries=priors, config=cfg,
            )

            story_task = tracer.traced_arun(
                story_agent,
                f"Write the opening story for chapter {brief.chapter_number}: {brief.title}",
                phase=f"StoryWriter-Ch{brief.chapter_number}",
            )
            content_task = tracer.traced_arun(
                content_agent,
                f"Write teaching content for chapter {brief.chapter_number}: {brief.title}",
                phase=f"ContentWriter-Ch{brief.chapter_number}",
            )

            story_resp, content_resp = await asyncio.gather(story_task, content_task)

            # Extract story markdown
            story_md = str(story_resp.content or "").strip()
            # Extract content markdown
            content_md = str(content_resp.content or "").strip()

            # Save intermediate parts
            write_text(parts_dir / f"ch_{brief.chapter_number:02d}_story.md", story_md)
            write_text(parts_dir / f"ch_{brief.chapter_number:02d}_content.md", content_md)

            story_wc = count_words(story_md)
            content_wc = count_words(content_md)
            print(f"      Ch {brief.chapter_number}: Story={story_wc}w, Content={content_wc}w")

            # ── Step 2: Combiner Writer (sequential) ───────────────────
            print(f"      Ch {brief.chapter_number}: Combiner merging...")

            combo_agent = combiner_writer_agent(
                state.config, state.blueprint, brief,
                story_markdown=story_md,
                content_markdown=content_md,
                prior_summaries=priors, config=cfg,
            )
            combo_resp = await tracer.traced_arun(
                combo_agent,
                f"Combine story and content for chapter {brief.chapter_number}: {brief.title}",
                phase=f"Combiner-Ch{brief.chapter_number}",
                output_schema=ChapterDraft,
            )

            if isinstance(combo_resp.content, ChapterDraft):
                draft = combo_resp.content
            else:
                raw = str(combo_resp.content or "").strip()
                draft = ChapterDraft(
                    chapter_number=brief.chapter_number,
                    title=brief.title,
                    content_markdown=raw,
                    word_count=count_words(raw),
                    summary=f"Chapter {brief.chapter_number} covering {brief.title}",
                )

            draft.chapter_number = brief.chapter_number
            draft.word_count = count_words(draft.content_markdown)
            state.drafts[brief.chapter_number] = draft
            state.rolling_summaries[brief.chapter_number] = draft.summary
            prior_summaries.append(f"Ch {brief.chapter_number} '{draft.title}': {draft.summary}")
            state.status["chapters_drafted"].append(brief.chapter_number)

            write_text(
                draft_dir / f"ch_{brief.chapter_number:02d}_{slugify(brief.title)}.md",
                draft.content_markdown,
            )
            print(f"      Ch {brief.chapter_number}: Final={draft.word_count}w")

    print(f"\n  Writing complete: {len(state.drafts)} chapters")
    state.save_snapshot()
    return state.drafts


# ── Phase 5: Editing (parallel) ────────────────────────────────────────────

async def run_editing(
    state: PipelineState,
    app_config: Optional[AppConfig] = None,
) -> Dict[int, EditedChapter]:
    """Edit all chapters — runs in parallel if configured."""
    cfg = app_config or load_config()
    tracer = state.tracer
    print("\n  Phase 5: Editing")
    print("  " + "-" * 50)
    state.status["phase"] = "editing"

    assert state.blueprint is not None, "Blueprint must be set before editing."
    edited_dir = ensure_dir(state.output_dir / "edited")

    async def edit_one(brief: ChapterBrief) -> EditedChapter:
        draft = state.drafts.get(brief.chapter_number)
        if draft is None:
            raise RuntimeError(f"No draft for chapter {brief.chapter_number}")

        state.status["current_step"] = f"editing_ch_{brief.chapter_number}"

        agent = editor_agent(state.config, brief, draft, config=cfg)
        prompt = f"Edit and polish chapter {brief.chapter_number}: {brief.title}"
        response = await tracer.traced_arun(
            agent, prompt,
            phase=f"Editor-Ch{brief.chapter_number}",
            output_schema=EditedChapter,
        )

        if isinstance(response.content, EditedChapter):
            edited = response.content
        else:
            raw = str(response.content or "").strip()
            edited = EditedChapter(
                chapter_number=brief.chapter_number,
                title=brief.title,
                content_markdown=raw or draft.content_markdown,
                final_word_count=count_words(raw or draft.content_markdown),
                changes_made=["Recovered from unstructured editor output"],
            )

        edited.chapter_number = brief.chapter_number
        edited.final_word_count = count_words(edited.content_markdown)
        state.status["chapters_edited"].append(brief.chapter_number)
        return edited

    if cfg.workflow.parallel_editing:
        tasks = [edit_one(brief) for brief in state.blueprint.chapters]
        results = await asyncio.gather(*tasks, return_exceptions=True)
    else:
        results = []
        for brief in state.blueprint.chapters:
            results.append(await edit_one(brief))

    for result in results:
        if isinstance(result, Exception):
            state.status["errors"].append(str(result))
            continue
        state.edited[result.chapter_number] = result
        write_text(
            edited_dir / f"ch_{result.chapter_number:02d}_{slugify(result.title)}.md",
            result.content_markdown,
        )

    print(f"\n  Editing complete: {len(state.edited)} chapters")
    state.save_snapshot()
    return state.edited


# ── Phase 5.5: QA Review ──────────────────────────────────────────────────

async def _rewrite_chapter(
    state: PipelineState,
    brief: ChapterBrief,
    revision_notes: str,
    cfg: AppConfig,
) -> EditedChapter:
    """Rewrite a single chapter through the full Story+Content+Combiner+Editor pipeline."""
    tracer = state.tracer
    research = state.research.get(brief.chapter_number)
    priors = list(state.rolling_summaries.values())
    ch_num = brief.chapter_number
    parts_dir = ensure_dir(state.output_dir / "drafts" / "parts")
    draft_dir = ensure_dir(state.output_dir / "drafts")
    edited_dir = ensure_dir(state.output_dir / "edited")

    # Inject revision notes into the brief synopsis so writers see the QA feedback
    enhanced_brief = brief.model_copy()

    # Step 1: Story + Content in parallel
    print(f"        Ch {ch_num}: Rewriting (Story + Content parallel)...")
    assigned_fmt = state.story_formats.get(ch_num)
    formats_so_far = [
        state.story_formats[n]
        for n in sorted(state.story_formats)
        if n < ch_num
    ]
    story_ag = story_writer_agent(
        state.config, state.blueprint, enhanced_brief,
        research=research, prior_summaries=priors,
        assigned_format=assigned_fmt,
        formats_used=formats_so_far,
        config=cfg,
    )
    content_ag = content_writer_agent(
        state.config, state.blueprint, enhanced_brief,
        research=research, prior_summaries=priors, config=cfg,
    )

    story_resp, content_resp = await asyncio.gather(
        tracer.traced_arun(
            story_ag,
            f"Rewrite opening story for ch {ch_num}. QA feedback: {revision_notes}",
            phase=f"QA-Rewrite-Story-Ch{ch_num}",
        ),
        tracer.traced_arun(
            content_ag,
            f"Rewrite teaching content for ch {ch_num}. QA feedback: {revision_notes}",
            phase=f"QA-Rewrite-Content-Ch{ch_num}",
        ),
    )

    story_md = str(story_resp.content or "").strip()
    content_md = str(content_resp.content or "").strip()
    write_text(parts_dir / f"ch_{ch_num:02d}_story_rewrite.md", story_md)
    write_text(parts_dir / f"ch_{ch_num:02d}_content_rewrite.md", content_md)

    # Step 2: Combiner
    print(f"        Ch {ch_num}: Combiner merging rewrite...")
    combo_ag = combiner_writer_agent(
        state.config, state.blueprint, enhanced_brief,
        story_markdown=story_md, content_markdown=content_md,
        prior_summaries=priors, config=cfg,
    )
    combo_resp = await tracer.traced_arun(
        combo_ag,
        f"Combine rewritten story+content for ch {ch_num}",
        phase=f"QA-Rewrite-Combine-Ch{ch_num}",
        output_schema=ChapterDraft,
    )

    if isinstance(combo_resp.content, ChapterDraft):
        draft = combo_resp.content
    else:
        raw = str(combo_resp.content or "").strip()
        draft = ChapterDraft(
            chapter_number=ch_num,
            title=brief.title,
            content_markdown=raw,
            word_count=count_words(raw),
            summary=f"Chapter {ch_num} rewritten after QA feedback",
        )
    draft.chapter_number = ch_num
    draft.word_count = count_words(draft.content_markdown)
    write_text(draft_dir / f"ch_{ch_num:02d}_{slugify(brief.title)}_rewrite.md", draft.content_markdown)

    # Step 3: Editor
    print(f"        Ch {ch_num}: Editing rewrite...")
    ed_ag = editor_agent(state.config, brief, draft, config=cfg)
    ed_resp = await tracer.traced_arun(
        ed_ag,
        f"Edit rewritten chapter {ch_num}",
        phase=f"QA-Rewrite-Edit-Ch{ch_num}",
        output_schema=EditedChapter,
    )

    if isinstance(ed_resp.content, EditedChapter):
        edited = ed_resp.content
    else:
        raw = str(ed_resp.content or "").strip()
        edited = EditedChapter(
            chapter_number=ch_num,
            title=brief.title,
            content_markdown=raw or draft.content_markdown,
            final_word_count=count_words(raw or draft.content_markdown),
            changes_made=["Rewritten after QA feedback"],
        )

    edited.chapter_number = ch_num
    edited.final_word_count = count_words(edited.content_markdown)
    write_text(edited_dir / f"ch_{ch_num:02d}_{slugify(brief.title)}.md", edited.content_markdown)
    print(f"        Ch {ch_num}: Rewrite complete ({edited.final_word_count}w)")
    return edited


async def run_qa_review(
    state: PipelineState,
    tools: Optional[List[Any]] = None,
    app_config: Optional[AppConfig] = None,
) -> None:
    """
    QA review with locked-pass protection and targeted recheck.

    Round 1: Full-book QA → lock passing chapters → rewrite failing ones.
    Round 2+: Targeted QA on ONLY rewritten chapters → locked chapters are never
              re-evaluated, preventing regression from LLM scoring inconsistency.
    """
    cfg = app_config or load_config()
    tracer = state.tracer
    max_retries = cfg.workflow.max_stage_retries

    print("\n  Phase 5.5: QA Review")
    print("  " + "-" * 50)
    state.status["phase"] = "qa_review"

    assert state.blueprint is not None
    assert state.edited, "No edited chapters to review."

    # revision_notes_pending tracks notes from the last round for rewritten chapters
    revision_notes_pending: Dict[int, str] = {}

    for attempt in range(max_retries + 1):
        chapters_sorted = sorted(state.edited.values(), key=lambda c: c.chapter_number)

        # ── Round 1: Full-book QA ─────────────────────────────────────────
        if attempt == 0:
            book_text = "\n\n---\n\n".join(
                f"## Chapter {ch.chapter_number}: {ch.title}\n\n{ch.content_markdown}"
                for ch in chapters_sorted
            )
            agent = qa_agent(
                state.config,
                all_chapters_markdown=book_text,
                tools=tools or [],
                config=cfg,
            )
            response = await tracer.traced_arun(
                agent,
                "Review the complete book for quality against Living Enlightenment standards.",
                phase="QA-Review-Round1",
                output_schema=BookQAReview,
            )

        # ── Round 2+: Targeted QA on rewritten chapters only ─────────────
        else:
            rewritten_nums = set(revision_notes_pending.keys())
            rewritten_text = "\n\n---\n\n".join(
                f"## Chapter {ch.chapter_number}: {ch.title}\n\n{ch.content_markdown}"
                for ch in chapters_sorted
                if ch.chapter_number in rewritten_nums
            )
            # Pass full book as context (truncated) so the agent can check continuity
            full_context = "\n\n---\n\n".join(
                f"## Chapter {ch.chapter_number}: {ch.title} [LOCKED — already passed]"
                for ch in chapters_sorted
                if ch.chapter_number in state.qa_locked
            )
            agent = targeted_qa_agent(
                state.config,
                rewritten_chapters_markdown=rewritten_text,
                full_book_context=full_context,
                revision_notes_by_chapter=revision_notes_pending,
                tools=tools or [],
                config=cfg,
            )
            response = await tracer.traced_arun(
                agent,
                f"Re-check {len(rewritten_nums)} rewritten chapter(s) after QA feedback.",
                phase=f"QA-Review-Round{attempt + 1}-Targeted",
                output_schema=BookQAReview,
            )

        if isinstance(response.content, BookQAReview):
            review = response.content
        else:
            print("    QA agent returned unstructured response — assuming PASS")
            return

        # Lock in any chapters that passed this round
        newly_locked = []
        for ch in review.chapters:
            if ch.passed and ch.chapter_number not in state.qa_locked:
                state.qa_locked.add(ch.chapter_number)
                newly_locked.append(ch.chapter_number)

        # Failing = only chapters that aren't locked
        failing = [ch for ch in review.chapters if not ch.passed and ch.chapter_number not in state.qa_locked]

        safe_print(f"\n    QA Round {attempt + 1} Results:")
        if attempt > 0 and state.qa_locked:
            locked_str = ", ".join(f"Ch{n}" for n in sorted(state.qa_locked))
            safe_print(f"      Locked (already passed): {locked_str}")
        for ch in review.chapters:
            if ch.chapter_number in state.qa_locked and ch.chapter_number not in newly_locked:
                continue  # don't re-print locked chapters
            status = "PASS ✓" if ch.passed else "FAIL"
            safe_print(f"      Ch {ch.chapter_number}: {status}  voice={ch.voice_score} structure={ch.structure_score} shastra={ch.shastra_score} story={ch.story_score}")
            if not ch.passed and ch.issues:
                for issue in ch.issues[:2]:
                    safe_print(f"        - {issue}")

        if review.book_level_notes:
            safe_print(f"    Book-level: {review.book_level_notes[:200]}")

        # Save this round's report
        write_json(
            state.output_dir / f"qa_review_round{attempt + 1}.json",
            review.model_dump(),
        )

        # All chapters either locked or passed — done
        total_chapters = len(state.edited)
        if not failing and len(state.qa_locked) >= total_chapters:
            safe_print(f"\n    QA APPROVED — all {total_chapters} chapters passed and locked")
            return

        if not failing:
            safe_print(f"\n    QA APPROVED — no failing chapters remain")
            return

        if attempt >= max_retries:
            safe_print(f"\n    QA: Max retries ({max_retries}) reached — proceeding with best available")
            safe_print(f"    Locked: {sorted(state.qa_locked)} | Still failing: {[ch.chapter_number for ch in failing]}")
            return

        # Rewrite only the failing (non-locked) chapters
        safe_print(f"\n    QA: {len(failing)} chapter(s) need rewrite — starting rewrite cycle...")
        revision_notes_pending = {}

        for ch_result in failing:
            brief = next(
                (b for b in state.blueprint.chapters if b.chapter_number == ch_result.chapter_number),
                None,
            )
            if brief is None:
                continue
            rewritten = await _rewrite_chapter(
                state, brief, ch_result.revision_notes, cfg,
            )
            state.edited[ch_result.chapter_number] = rewritten
            revision_notes_pending[ch_result.chapter_number] = ch_result.revision_notes

        state.save_snapshot()
        safe_print(f"    Rewrite cycle complete — running targeted QA on {len(revision_notes_pending)} chapter(s)...")


# ── Section heading removal ────────────────────────────────────────────────

_SECTION_HEADINGS = re.compile(
    r"^###\s*(Opening Story|The Teaching|Practical Exercise|Humor|Closing Bridge|Transition)\s*$",
    re.MULTILINE,
)


def _strip_section_headings(markdown: str) -> str:
    """Remove internal section headings (### Opening Story, etc.) from final output.
    These are structural markers used during the pipeline but should not appear in the book."""
    return _SECTION_HEADINGS.sub("", markdown)


# ── Phase 5.55: Q&A Generation (parallel per chapter) ────────────────────

async def run_qa_generation(
    state: PipelineState,
    tools: Optional[List[Any]] = None,
    app_config: Optional[AppConfig] = None,
) -> None:
    """Generate Q&A pairs for every chapter in parallel after QA review."""
    cfg = app_config or load_config()
    tracer = state.tracer

    if not cfg.workflow.include_chapter_qa:
        return

    print("\n  Phase 5.55: Q&A Generation")
    print("  " + "-" * 50)
    state.status["phase"] = "qa_generation"

    assert state.blueprint is not None
    assert state.edited, "No edited chapters for Q&A generation."

    mcp_conn: Optional[McpConnection] = getattr(state, "_mcp_conn", None)

    async def generate_qa_one(brief: ChapterBrief) -> Optional[ChapterQA]:
        edited_ch = state.edited.get(brief.chapter_number)
        if not edited_ch:
            return None

        # Each Q&A agent gets its own MCP pool
        agent_tools: Optional[List[Any]] = None
        if mcp_conn and cfg:
            agent_tools = await mcp_conn.get_pool_for_agent(cfg)
        elif tools:
            agent_tools = tools

        agent = qa_generator_agent(
            state.config, brief,
            chapter_content=edited_ch.content_markdown,
            tools=agent_tools or [],
            config=cfg,
        )
        response = await tracer.traced_arun(
            agent,
            f"Generate Q&A pairs for chapter {brief.chapter_number}: {brief.title}",
            phase=f"QAGen-Ch{brief.chapter_number}",
            output_schema=ChapterQA,
        )

        if isinstance(response.content, ChapterQA):
            return response.content

        # Fallback: try to parse raw text
        raw = str(response.content or "").strip()
        if raw:
            # Build a minimal ChapterQA from raw text if structured output failed
            return ChapterQA(
                chapter_number=brief.chapter_number,
                chapter_title=brief.title,
                qa_pairs=[QAPair(
                    question="What is the core teaching of this chapter?",
                    answer=raw[:500],
                    source_reference="",
                )],
            )
        return None

    tasks = [generate_qa_one(brief) for brief in state.blueprint.chapters]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    qa_dir = ensure_dir(state.output_dir / "chapter_qa")
    total_pairs = 0

    for result in results:
        if isinstance(result, Exception):
            state.status["errors"].append(f"Q&A generation error: {result}")
            continue
        if result is None:
            continue

        state.chapter_qa[result.chapter_number] = result
        total_pairs += len(result.qa_pairs)

        # Append Q&A markdown to the edited chapter content
        edited_ch = state.edited.get(result.chapter_number)
        if edited_ch:
            qa_md = _render_qa_markdown(result)
            edited_ch.content_markdown = edited_ch.content_markdown.rstrip() + "\n\n" + qa_md
            edited_ch.final_word_count = count_words(edited_ch.content_markdown)

        write_json(
            qa_dir / f"ch_{result.chapter_number:02d}_qa.json",
            result.model_dump(),
        )

    print(f"  Q&A generation complete: {len(state.chapter_qa)} chapters, {total_pairs} total pairs")
    state.save_snapshot()


def _render_qa_markdown(chapter_qa: ChapterQA) -> str:
    """Render a ChapterQA as a markdown section appended to the chapter."""
    lines = [
        "---",
        "",
        "## Questions and Answers",
        "",
    ]
    for i, pair in enumerate(chapter_qa.qa_pairs, 1):
        lines.append(f"**Q{i}: {pair.question}**")
        lines.append("")
        lines.append(pair.answer)
        if pair.source_reference:
            lines.append(f"")
            lines.append(f"*— {pair.source_reference}*")
        lines.append("")
    return "\n".join(lines)


# ── Phase 5.6: Frontmatter (Foreword + Benediction) ──────────────────────

async def run_frontmatter(
    state: PipelineState,
    app_config: Optional[AppConfig] = None,
) -> None:
    """Generate foreword and benediction in parallel after QA is complete."""
    cfg = app_config or load_config()
    tracer = state.tracer

    needs_foreword = state.config.include_foreword
    needs_benediction = state.config.include_benediction

    if not needs_foreword and not needs_benediction:
        return

    print("\n  Phase 5.6: Frontmatter (Foreword + Benediction)")
    print("  " + "-" * 50)
    state.status["phase"] = "frontmatter"

    assert state.blueprint is not None

    # Extract the last chapter's closing bridge to give the benediction a natural lead-in
    last_chapter_bridge = ""
    if state.edited:
        last_ch_num = max(state.edited.keys())
        last_content = state.edited[last_ch_num].content_markdown
        bridge_match = re.search(
            r"###\s*Closing Bridge\s*(.*?)(?=###|\Z)", last_content, re.DOTALL | re.IGNORECASE
        )
        if bridge_match:
            last_chapter_bridge = bridge_match.group(1).strip()[:500]

    tasks = []
    task_labels = []

    if needs_foreword:
        fw_ag = foreword_agent(state.config, state.blueprint, cfg)
        tasks.append(tracer.traced_arun(
            fw_ag,
            f"Write the foreword for '{state.config.title}'.",
            phase="Frontmatter-Foreword",
        ))
        task_labels.append("foreword")

    if needs_benediction:
        bene_ag = benediction_agent(state.config, state.blueprint, last_chapter_bridge, cfg)
        tasks.append(tracer.traced_arun(
            bene_ag,
            f"Write the closing benediction for '{state.config.title}'.",
            phase="Frontmatter-Benediction",
        ))
        task_labels.append("benediction")

    results = await asyncio.gather(*tasks, return_exceptions=True)

    foreword_text = ""
    benediction_text = ""

    for label, result in zip(task_labels, results):
        if isinstance(result, Exception):
            state.status["errors"].append(f"Frontmatter {label} failed: {result}")
            print(f"  WARNING: {label} generation failed — will be omitted from output")
            continue
        text = str(result.content or "").strip()
        wc = count_words(text)
        if label == "foreword":
            foreword_text = text
            print(f"  Foreword: {wc} words")
            write_text(state.output_dir / "foreword.md", text)
        else:
            benediction_text = text
            print(f"  Benediction: {wc} words")
            write_text(state.output_dir / "benediction.md", text)

    # Attach to state so run_designer can pick them up
    state._foreword = foreword_text
    state._benediction = benediction_text


# ── Phase 6: Design (.docx) ───────────────────────────────────────────────

async def run_designer(
    state: PipelineState,
    app_config: Optional[AppConfig] = None,
) -> Path:
    """Assemble final .docx from edited chapters."""
    cfg = app_config or load_config()
    print("\n  Phase 6: Designer")
    print("  " + "-" * 50)
    state.status["phase"] = "design"
    state.status["current_step"] = "assembling_docx"

    assert state.blueprint is not None
    assert state.edited, "No edited chapters to assemble."

    # Strip internal section headings from all chapters before final output
    for ch_num, ch in state.edited.items():
        ch.content_markdown = _strip_section_headings(ch.content_markdown)

    chapters_sorted = sorted(state.edited.values(), key=lambda c: c.chapter_number)
    total_words = sum(c.final_word_count for c in chapters_sorted)

    state.metadata = CompiledBookMetadata(
        title=state.config.title,
        subtitle=state.config.subtitle,
        author=state.config.author,
        synopsis=state.config.synopsis,
        foreword=getattr(state, "_foreword", ""),
        benediction=getattr(state, "_benediction", ""),
        source_notes=[s for s in state.config.reference_sources] if state.config.reference_sources else [],
        all_source_links=state.all_source_links,
        chapter_count=len(chapters_sorted),
        estimated_word_count=total_words,
    )

    docx_path = state.output_dir / f"{slugify(state.config.title)}.docx"
    _build_docx(state, chapters_sorted, docx_path)

    md_path = state.output_dir / f"{slugify(state.config.title)}.md"
    _build_combined_markdown(state, chapters_sorted, md_path)

    state.docx_path = docx_path
    print(f"  DOCX saved: {docx_path}")
    print(f"  Markdown saved: {md_path}")
    state.save_snapshot()
    return docx_path


def _build_docx(state: PipelineState, chapters: List[EditedChapter], output_path: Path) -> None:
    """Build a styled .docx file from edited chapters."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("  WARNING: python-docx not installed. Skipping .docx generation.")
        print("  Install with: pip install python-docx")
        return

    doc = Document()
    config = state.config

    style = doc.styles["Normal"]
    style.font.name = "Palatino Linotype"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)

    for level, (name_suffix, size, bold, color) in {
        0: ("Title", 28, True, RGBColor(0x1A, 0x1A, 0x2E)),
        1: ("Heading 1", 22, True, RGBColor(0x1A, 0x1A, 0x2E)),
        2: ("Heading 2", 14, True, RGBColor(0x0F, 0x34, 0x60)),
        3: ("Heading 3", 12, True, RGBColor(0x0F, 0x34, 0x60)),
    }.items():
        try:
            h_style = doc.styles[name_suffix]
        except KeyError:
            h_style = doc.styles.add_style(name_suffix, WD_STYLE_TYPE.PARAGRAPH)
        h_style.font.name = "Georgia"
        h_style.font.size = Pt(size)
        h_style.font.bold = bold
        h_style.font.color.rgb = color

    try:
        quote_style = doc.styles["Quote"]
    except KeyError:
        quote_style = doc.styles.add_style("Quote", WD_STYLE_TYPE.PARAGRAPH)
    quote_style.font.name = "Palatino Linotype"
    quote_style.font.size = Pt(11)
    quote_style.font.italic = True

    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run(config.title)
    title_run.font.name = "Georgia"
    title_run.font.size = Pt(28)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    if config.subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sub_run = sub_para.add_run(config.subtitle)
        sub_run.font.name = "Georgia"
        sub_run.font.size = Pt(16)
        sub_run.italic = True
        sub_run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)

    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_run = author_para.add_run(f"\n\n{config.author}")
    author_run.font.name = "Georgia"
    author_run.font.size = Pt(14)

    doc.add_page_break()

    if config.include_toc:
        doc.add_heading("Table of Contents", level=1)
        for ch in chapters:
            doc.add_paragraph(
                f"Chapter {ch.chapter_number}: {ch.title}",
                style="Normal",
            )
        doc.add_page_break()

    if config.include_foreword and state.metadata and state.metadata.foreword:
        doc.add_heading("Foreword", level=1)
        doc.add_paragraph(state.metadata.foreword)
        doc.add_page_break()

    # Build a map of section_group → first chapter_number so we know when to insert dividers
    cfg = load_config()
    section_dividers_enabled = cfg.workflow.include_section_dividers
    seen_groups: set = set()
    # Map chapter_number → section_group from blueprint
    chapter_groups: Dict[int, str] = {}
    group_descriptions: Dict[str, str] = {}
    if state.blueprint:
        for ch in state.blueprint.chapters:
            chapter_groups[ch.chapter_number] = ch.section_group or ""
        for sec in (state.blueprint.sections or []):
            group_descriptions[sec.label] = sec.description

    for ch in chapters:
        grp = chapter_groups.get(ch.chapter_number, "")

        # Insert section divider page when a new named group starts
        if section_dividers_enabled and grp and grp not in seen_groups:
            seen_groups.add(grp)
            # Full-page section divider: centered label + description
            sec_para = doc.add_paragraph()
            sec_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Add several blank lines to push text to vertical center
            sec_para.add_run("\n" * 8)
            sec_run = sec_para.add_run(grp)
            sec_run.font.name = "Georgia"
            sec_run.font.size = Pt(28)
            sec_run.bold = True
            sec_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            desc = group_descriptions.get(grp, "")
            if desc:
                desc_para = doc.add_paragraph()
                desc_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                desc_run = desc_para.add_run(desc)
                desc_run.font.name = "Palatino Linotype"
                desc_run.font.size = Pt(12)
                desc_run.italic = True
                desc_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            doc.add_page_break()

        doc.add_heading(f"Chapter {ch.chapter_number}: {ch.title}", level=1)
        _markdown_to_docx(doc, ch.content_markdown)
        doc.add_page_break()

    if config.include_benediction and state.metadata and state.metadata.benediction:
        doc.add_heading("Benediction", level=1)
        doc.add_paragraph(state.metadata.benediction)

    # Source links appendix
    if state.metadata and state.metadata.all_source_links:
        doc.add_page_break()
        doc.add_heading("Sources", level=1)
        doc.add_paragraph(
            "The following satsangs and teachings were referenced during the creation of this book.",
            style="Normal",
        )
        current_ch = None
        for lnk in sorted(state.metadata.all_source_links, key=lambda x: x.chapter_number):
            if lnk.chapter_number != current_ch:
                current_ch = lnk.chapter_number
                ch_title = next(
                    (c.title for c in chapters if c.chapter_number == current_ch), f"Chapter {current_ch}"
                )
                doc.add_heading(f"Chapter {current_ch}: {ch_title}", level=2)
            entry = f"{lnk.title}"
            if lnk.date:
                entry += f" — {lnk.date}"
            entry += f"\n{lnk.url}"
            doc.add_paragraph(entry, style="Normal")

    doc.save(str(output_path))


def _markdown_to_docx(doc, markdown: str) -> None:
    """Simple markdown-to-docx converter for chapter content."""
    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            pass
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("> "):
            doc.add_paragraph(stripped[2:], style="Quote")
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(stripped, style="List Number")
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            para = doc.add_paragraph()
            _add_formatted_runs(para, stripped)

        i += 1


def _add_formatted_runs(paragraph, text: str) -> None:
    """Parse inline markdown formatting (bold, italic) into docx runs."""
    parts = re.split(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _build_combined_markdown(state: PipelineState, chapters: List[EditedChapter], output_path: Path) -> None:
    """Build a combined markdown file as a reference copy."""
    config = state.config
    parts = [f"# {config.title}"]
    if config.subtitle:
        parts.append(f"### {config.subtitle}")
    parts.append(f"\n*{config.author}*\n")

    # TOC with section groupings
    if config.include_toc:
        parts.append("---\n\n## Table of Contents\n")
        seen_groups: set = set()
        chapter_groups: Dict[int, str] = {}
        if state.blueprint:
            for ch in state.blueprint.chapters:
                chapter_groups[ch.chapter_number] = ch.section_group or ""
        for ch in chapters:
            grp = chapter_groups.get(ch.chapter_number, "")
            if grp and grp not in seen_groups:
                seen_groups.add(grp)
                parts.append(f"\n**{grp}**\n")
            anchor = f"chapter-{ch.chapter_number}-{slugify(ch.title).lower()}"
            parts.append(f"- [Chapter {ch.chapter_number}: {ch.title}](#{anchor})")
        parts.append("")

    if config.include_foreword and state.metadata and state.metadata.foreword:
        parts.append("---\n\n## Foreword\n")
        parts.append(state.metadata.foreword)

    # Chapters with section dividers
    seen_groups_md: set = set()
    chapter_groups_md: Dict[int, str] = {}
    if state.blueprint:
        for ch in state.blueprint.chapters:
            chapter_groups_md[ch.chapter_number] = ch.section_group or ""

    parts.append("\n---\n")
    for ch in chapters:
        grp = chapter_groups_md.get(ch.chapter_number, "")
        if grp and grp not in seen_groups_md:
            seen_groups_md.add(grp)
            parts.append(f"\n---\n\n# {grp}\n\n---\n")
        parts.append(ch.content_markdown.strip())
        parts.append("\n---\n")

    if config.include_benediction and state.metadata and state.metadata.benediction:
        parts.append("## Benediction\n")
        parts.append(state.metadata.benediction)

    # Sources appendix
    if state.metadata and state.metadata.all_source_links:
        parts.append("\n---\n\n## Sources\n")
        parts.append("The following satsangs and teachings were referenced in this book.\n")
        current_ch = None
        for lnk in sorted(state.metadata.all_source_links, key=lambda x: x.chapter_number):
            if lnk.chapter_number != current_ch:
                current_ch = lnk.chapter_number
                ch_title = next(
                    (c.title for c in chapters if c.chapter_number == current_ch),
                    f"Chapter {current_ch}",
                )
                parts.append(f"\n### Chapter {current_ch}: {ch_title}\n")
            date_str = f" — {lnk.date}" if lnk.date else ""
            parts.append(f"- **{lnk.title}**{date_str}  \n  {lnk.url}")

    write_text(output_path, "\n".join(parts))


# ── Main Pipeline ──────────────────────────────────────────────────────────

async def run_pipeline(
    raw_input: dict,
    mcp_tools: Optional[List[Any]] = None,
    config_path: Optional[str] = None,
    skip_mcp: bool = False,
) -> PipelineState:
    """
    Run the full book generation pipeline.

    Args:
        raw_input: User-provided book configuration (can be partial).
        mcp_tools: Optional pre-connected MCP tool list (overrides auto-connect).
        config_path: Optional path to config.yaml override.
        skip_mcp: If True, skip MCP connections entirely (agents use training data only).

    Returns:
        PipelineState with all artifacts.
    """
    from config import reload_config
    app_config = reload_config(config_path)

    # ── Phase 1: Intake (tracer not yet output-bound) ──────────────────
    temp_tracer = Tracer(output_dir=None, verbose=True)
    book_config = await run_intake(raw_input, app_config, temp_tracer)

    # ── Setup output directory ─────────────────────────────────────────
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_root = Path(app_config.workflow.output_dir)
    output_dir = ensure_dir(output_root / f"{slugify(book_config.title)}_{timestamp}")
    write_json(output_dir / "book_config.json", book_config.model_dump())

    tracer = Tracer(output_dir=output_dir, verbose=True)
    tracer.traces.extend(temp_tracer.traces)

    state = PipelineState(book_config, output_dir, tracer)

    # ── Connect to MCP servers ─────────────────────────────────────────
    mcp_conn = McpConnection()
    state._mcp_conn = mcp_conn  # store for parallel research pools
    tools = mcp_tools  # use provided tools if any
    if tools is None and not skip_mcp and app_config.mcp_sources:
        print("\n  Connecting to MCP sources...")
        print("  " + "-" * 50)
        mcp_list = await mcp_conn.connect(app_config)
        if mcp_list:
            tools = mcp_list
            print(f"  {len(mcp_list)} MCP source(s) connected")
        else:
            print("  No MCP sources connected — agents will use training data only")

    print("\n" + "=" * 72)
    print(f"  Book Generator v2")
    print(f"  Title: {book_config.title}")
    print(f"  Chapters: {book_config.num_chapters}")
    print(f"  Words/chapter: {book_config.words_per_chapter}")
    print(f"  POV: {book_config.pov} | Tone: {book_config.tone}")
    print(f"  MCP tools: {len(tools) if tools else 0}")
    print(f"  Output: {output_dir}")
    print(f"  Traces: {output_dir / 'agent_traces.log'}")
    print("=" * 72)

    try:
        # ── Phase 2: Architect ─────────────────────────────────────────
        await run_architect(state, tools=tools, app_config=app_config)

        # ── Phase 3: Research (parallel) ───────────────────────────────
        await run_research(state, tools=tools, app_config=app_config)

        # ── Phase 3.5: YouTube Video Matching ──────────────────────────
        await run_youtube_matching(state, app_config=app_config)

        # ── Phase 4: Writing (batched) ─────────────────────────────────
        await run_writing(state, app_config=app_config)

        # ── Phase 5: Editing (parallel) ────────────────────────────────
        await run_editing(state, app_config=app_config)

        # ── Phase 5.5: QA Review (with rewrite loop) ─────────────────
        await run_qa_review(state, tools=tools, app_config=app_config)

        # ── Phase 5.55: Q&A Generation (parallel per chapter) ─────────
        await run_qa_generation(state, tools=tools, app_config=app_config)

        # ── Phase 5.6: Frontmatter (Foreword + Benediction) ───────────
        await run_frontmatter(state, app_config=app_config)

        # ── Phase 6: Designer (.docx) ──────────────────────────────────
        await run_designer(state, app_config=app_config)

    finally:
        # Always close MCP connections
        await mcp_conn.close()

    # ── Done ───────────────────────────────────────────────────────────
    state.status["phase"] = "complete"
    state.status["current_step"] = "done"
    state.save_snapshot("final_state")

    print(tracer.summary())

    print("=" * 72)
    print("  Book generation complete!")
    print(f"  DOCX: {state.docx_path}")
    print(f"  Chapters: {len(state.edited)}")
    total = sum(c.final_word_count for c in state.edited.values())
    print(f"  Total words: {total}")
    print(f"  Output dir: {output_dir}")
    print("=" * 72 + "\n")

    return state
