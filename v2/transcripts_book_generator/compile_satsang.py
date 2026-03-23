"""
Satsang Compilation Pipeline — compiles transcript files into a book.

Unlike the creative book generator, this pipeline PRESERVES every word
of the original transcripts. Agents only add structure, formatting,
and supplementary shastra pramanas.

Usage:
    python compile_satsang.py --dir transcripts_feb2026
    python compile_satsang.py --dir /path/to/transcripts
"""

from __future__ import annotations

import asyncio
import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional
from textwrap import dedent

from pydantic import BaseModel, Field

from config import AppConfig, build_agent_model, load_config, reload_config
from tracing import Tracer
from agents import SWAMIJI_VOICE

from agno.agent import Agent
from agno.tools.mcp import MCPTools


# ── Compilation Models ─────────────────────────────────────────────────────

class TranscriptFile(BaseModel):
    """One ingested transcript."""
    day_number: int
    date: str
    title: str
    filename: str
    content: str
    word_count: int


class ChapterPlan(BaseModel):
    """Plan for one chapter (one satsang day)."""
    chapter_number: int
    day_number: int
    date: str
    chapter_title: str
    occasion: str = ""
    key_themes: List[str] = Field(default_factory=list)
    existing_shastra_refs: List[str] = Field(default_factory=list)
    suggested_additional_refs: List[str] = Field(default_factory=list)


class BookStructure(BaseModel):
    """Overall book structure from the Structurer agent."""
    book_title: str
    subtitle: str = ""
    overview: str = ""
    chapters: List[ChapterPlan]


class ShastraEnrichment(BaseModel):
    """Additional shastra pramanas for one chapter."""
    chapter_number: int
    pramanas: List[dict] = Field(
        default_factory=list,
        description='List of {"sanskrit": ..., "source": ..., "translation": ..., "insert_after": ...}',
    )


class FormattedChapter(BaseModel):
    """A formatted chapter ready for assembly."""
    chapter_number: int
    title: str
    date: str
    content_markdown: str
    word_count: int
    original_word_count: int


class CompilationQAResult(BaseModel):
    """QA result for the compilation."""
    approved: bool
    chapters_checked: int
    words_original: int
    words_final: int
    issues: List[str] = Field(default_factory=list)
    notes: str = ""


# ── Utilities ──────────────────────────────────────────────────────────────

def safe_print(text: str) -> None:
    try:
        print(text)
    except UnicodeEncodeError:
        print(text.encode("ascii", "replace").decode())


def slugify(value: str) -> str:
    cleaned = "".join(ch if ch.isalnum() or ch in (" ", "-", "_") else "_" for ch in value)
    return "_".join(cleaned.split())[:80] or "book"


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def write_json(path: Path, data: Any) -> None:
    ensure_dir(path.parent)
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


def write_text(path: Path, content: str) -> None:
    ensure_dir(path.parent)
    path.write_text(content, encoding="utf-8")


def count_words(text: str) -> int:
    return len(text.split())


# ── Phase 1: Ingest ───────────────────────────────────────────────────────

def ingest_transcripts(transcript_dir: Path) -> List[TranscriptFile]:
    """Read all transcript files, parse day numbers, sort chronologically."""
    print("\n  Phase 1: Ingest")
    print("  " + "-" * 50)

    files = sorted(transcript_dir.glob("*.md")) + sorted(transcript_dir.glob("*.txt"))
    if not files:
        raise FileNotFoundError(f"No transcript files found in {transcript_dir}")

    transcripts = []
    for fpath in files:
        content = fpath.read_text(encoding="utf-8")
        fname = fpath.name

        # Parse day number from filename
        day_match = re.search(r'[Dd]ay\s*(\d+)', fname)
        day_num = int(day_match.group(1)) if day_match else 0

        # Parse date from filename
        date_match = re.search(r'(\d{1,2})\s*(Feb|Jan|Mar)\s*(\d{4})', fname)
        if date_match:
            day_d, month, year = date_match.groups()
            month_num = {"Jan": "01", "Feb": "02", "Mar": "03"}.get(month, "01")
            date_str = f"{year}-{month_num}-{day_d.zfill(2)}"
        else:
            date_str = "unknown"

        # Parse title from content (first SATSANG TITLE line)
        title_match = re.search(r'SATSANG TITLE\s*\|\s*(.+)', content)
        title = title_match.group(1).strip() if title_match else fname

        wc = count_words(content)
        transcripts.append(TranscriptFile(
            day_number=day_num,
            date=date_str,
            title=title,
            filename=fname,
            content=content,
            word_count=wc,
        ))
        safe_print(f"  Day {day_num:>2} | {date_str} | {wc:>6} words | {fname[:60]}")

    # Sort by day number
    transcripts.sort(key=lambda t: t.day_number)

    total = sum(t.word_count for t in transcripts)
    print(f"\n  Total: {len(transcripts)} transcripts, {total:,} words")
    return transcripts


# ── Phase 1.5: Content Cleanup (deterministic) ────────────────────────────

class CleanupReport(BaseModel):
    """Report of all changes made during transcript cleanup."""
    chapter_number: int = 0
    repetitions_removed: List[str] = Field(default_factory=list)
    half_phrases_removed: List[str] = Field(default_factory=list)
    discourse_markers_thinned: List[str] = Field(default_factory=list)
    vip_welcomes_removed: List[str] = Field(default_factory=list)
    original_word_count: int = 0
    cleaned_word_count: int = 0


def _remove_verbatim_repetitions(text: str) -> tuple:
    """
    Remove back-to-back duplicate sentences/phrases.
    Keep non-English blessings and mantras even if repeated.
    """
    removed = []

    # Split into sentences
    sentences = re.split(r'(?<=[.!?])\s+', text)
    if len(sentences) < 2:
        return text, removed

    cleaned = [sentences[0]]
    for i in range(1, len(sentences)):
        curr = sentences[i].strip()
        prev = sentences[i - 1].strip()

        # Normalize for comparison (lowercase, strip punctuation)
        curr_norm = re.sub(r'[^\w\s]', '', curr.lower()).strip()
        prev_norm = re.sub(r'[^\w\s]', '', prev.lower()).strip()

        if curr_norm == prev_norm and curr_norm:
            # Check if it's a mantra/blessing (Sanskrit/non-English) — keep those
            sanskrit_markers = ['om', 'namah', 'paramashiv', 'praptirastu', 'mangala', 'astu', 'svaha', 'nithyananda']
            is_mantra = any(m in curr_norm for m in sanskrit_markers) and len(curr.split()) < 15
            if is_mantra:
                cleaned.append(curr)
            else:
                removed.append(curr[:80] + "..." if len(curr) > 80 else curr)
        else:
            cleaned.append(curr)

    return " ".join(cleaned), removed


def _remove_half_phrases(text: str) -> tuple:
    """
    Remove false starts where Swamiji begins a phrase, stops, and restarts.
    Pattern: "I really... When you said white, I really saw white" → remove "I really..."
    """
    removed = []

    # Pattern: short fragment ending with "..." followed by a longer sentence starting similarly
    pattern = re.compile(
        r'(\b\w+(?:\s+\w+){0,4})\s*\.{2,3}\s+'  # short fragment + ellipsis
        r'(\1\b)',  # same words restart
        re.IGNORECASE,
    )

    def _replacer(match):
        fragment = match.group(1) + "..."
        removed.append(fragment)
        return match.group(2)

    cleaned = pattern.sub(_replacer, text)

    # Also catch: "word word... word word word" pattern (incomplete thought restarted)
    # More aggressive: remove "X Y... " when followed by "X Y Z" within 50 chars
    pattern2 = re.compile(r'(\b\w+\s+\w+)\s*\.{2,3}\s+(?=\1)', re.IGNORECASE)
    cleaned = pattern2.sub(lambda m: (removed.append(m.group(0).strip()), '')[1] or '', cleaned)

    return cleaned, removed


def _thin_discourse_markers(text: str) -> tuple:
    """
    Remove excessive standalone discourse markers (Listen, Understand, I tell you)
    when they appear too close together. Keep those before important truths.
    """
    removed = []

    # Split into paragraphs (bullet points)
    lines = text.split('\n')
    cleaned_lines = []
    last_marker_idx = -999  # track distance between markers

    marker_pattern = re.compile(
        r'^\s*\*?\s*(Listen\.?|Understand\.?|Understand deeply\.?|I tell you\.?|'
        r'Now,?\s*listen\.?|Listen\s+intensely\.?|Listen\s+carefully\.?)\s*$',
        re.IGNORECASE,
    )

    for i, line in enumerate(lines):
        if marker_pattern.match(line.strip()):
            distance = i - last_marker_idx
            if distance < 5:  # too close — remove
                removed.append(line.strip()[:60])
                continue
            last_marker_idx = i
        cleaned_lines.append(line)

    return '\n'.join(cleaned_lines), removed


def _remove_vip_welcomes(text: str) -> tuple:
    """
    Remove sections where Bhagavan is welcoming VIPs, guests from various countries.
    These are not needed for the book.
    """
    removed = []

    # Pattern: lines mentioning welcoming VIPs, presidents, prime ministers, countries
    vip_pattern = re.compile(
        r'(?:^|\n)\s*\*?\s*(?:I WELCOME|ALL THE (?:IMPORTANT )?VIPS?|'
        r'STATESMEN|PRESIDENTS|PRIME MINISTERS|FROM ALL OVER THE WORLD|'
        r'SARVAJNAPEETHA|CHIT SABHA|RAJA SABHA|RAJYA SABHA MEMBERS|'
        r'ALL THOSE WHO HAVE GATHERED|KAILASAVAASIS|'
        r'(?:I )?WELCOME (?:ALL|EACH|EVERY)(?:ONE)?(?:\s+OF YOU)?(?:\s+FOR BEING HERE)?)'
        r'[^.!?\n]*[.!?\n]',
        re.IGNORECASE | re.MULTILINE,
    )

    matches = vip_pattern.findall(text)
    for m in matches:
        stripped = m.strip()
        # Don't remove the very first welcome ("I welcome all of you with my love and blessings")
        if 'love and blessings' in stripped.lower() or 'love and respects' in stripped.lower():
            continue
        removed.append(stripped[:80] + "..." if len(stripped) > 80 else stripped)

    cleaned = vip_pattern.sub(lambda m: '' if 'love and blessings' not in m.group().lower() and 'love and respects' not in m.group().lower() else m.group(), text)

    return cleaned, removed


def run_content_cleanup(
    transcripts: List[TranscriptFile],
) -> tuple:
    """
    Deterministic content cleanup — no LLM.
    Returns cleaned transcripts and a report of all changes.
    """
    print("\n  Phase 1.5: Content Cleanup")
    print("  " + "-" * 50)

    cleaned_transcripts = []
    all_reports = []

    for t in transcripts:
        content = t.content
        report = CleanupReport(
            chapter_number=t.day_number,
            original_word_count=t.word_count,
        )

        # Step 1: Remove VIP welcomes
        content, vip_removed = _remove_vip_welcomes(content)
        report.vip_welcomes_removed = vip_removed

        # Step 2: Remove verbatim repetitions
        content, reps_removed = _remove_verbatim_repetitions(content)
        report.repetitions_removed = reps_removed

        # Step 3: Remove half phrases / false starts
        content, halfs_removed = _remove_half_phrases(content)
        report.half_phrases_removed = halfs_removed

        # Step 4: Thin discourse markers
        content, markers_removed = _thin_discourse_markers(content)
        report.discourse_markers_thinned = markers_removed

        report.cleaned_word_count = count_words(content)
        all_reports.append(report)

        cleaned = TranscriptFile(
            day_number=t.day_number,
            date=t.date,
            title=t.title,
            filename=t.filename,
            content=content,
            word_count=count_words(content),
        )
        cleaned_transcripts.append(cleaned)

        diff = cleaned.word_count - t.word_count
        removals = len(vip_removed) + len(reps_removed) + len(halfs_removed) + len(markers_removed)
        safe_print(
            f"  Day {t.day_number:>2}: {t.word_count}w -> {cleaned.word_count}w "
            f"({diff:+d}w, {removals} removals)"
        )

    total_orig = sum(t.word_count for t in transcripts)
    total_clean = sum(t.word_count for t in cleaned_transcripts)
    safe_print(f"\n  Total: {total_orig:,}w -> {total_clean:,}w ({total_clean - total_orig:+,}w)")

    return cleaned_transcripts, all_reports


# ── Phase 2: Structure ─────────────────────────────────────────────────────

async def run_structurer(
    transcripts: List[TranscriptFile],
    tracer: Tracer,
    config: AppConfig,
) -> BookStructure:
    """Create book structure from transcripts."""
    print("\n  Phase 2: Structure")
    print("  " + "-" * 50)

    # Build a compact summary of each transcript for the structurer
    summaries = []
    for t in transcripts:
        # First 800 chars + last 300 chars to capture key themes
        preview = t.content[:800] + "\n...\n" + t.content[-300:]
        summaries.append(
            f"DAY {t.day_number} ({t.date}): {t.title}\n"
            f"Word count: {t.word_count}\n"
            f"Preview:\n{preview}\n"
        )

    agent = Agent(
        name="Structurer",
        role="Plan the book structure from satsang transcripts.",
        model=build_agent_model("architect", config),
        instructions=[
            "You are structuring a book compiled from satsang transcripts.",
            "Each transcript becomes ONE chapter — do NOT merge or split transcripts.",
            "",
            f"There are {len(transcripts)} transcripts (Days: {', '.join(str(t.day_number) for t in transcripts)}).",
            "",
            "For each chapter provide:",
            "  - chapter_title: MUST capture the SPECIFIC teaching/action given to the reader,",
            "    not just the theme. BAD: 'Entangled with the Divine'. GOOD: 'Simply Assume What You",
            "    Want to Be: The First Quantum Leap'. The title should tell the reader what THEY",
            "    will learn to DO in this chapter.",
            "  - occasion: the occasion (e.g. 'Maha Shivaratri Brahmotsavam Day 3')",
            "  - key_themes: 3-5 key themes covered in that day's satsang",
            "  - existing_shastra_refs: any Sanskrit verses/shlokas already in the transcript",
            "  - suggested_additional_refs: 8-10 shastra references that could supplement the teaching",
            "    (we need at least 10 pramanas per chapter from Vedas, Agamas, Upanishads, Gita, Yoga Sutras)",
            "",
            "Also provide:",
            "  - book_title: overall title including 'Quantum Entanglement Quantum Enlightenment'",
            "  - subtitle: a descriptive subtitle",
            "  - overview: 2-3 sentence overview of the series",
            "",
            "TRANSCRIPT SUMMARIES:",
            "\n---\n".join(summaries),
        ],
        markdown=False,
    )

    response = await tracer.traced_arun(
        agent,
        "Create the book structure from these satsang transcripts.",
        phase="Structurer",
        output_schema=BookStructure,
    )

    if isinstance(response.content, BookStructure):
        structure = response.content
    else:
        # Fallback: create basic structure
        structure = BookStructure(
            book_title="Quantum Entanglement, Quantum Enlightenment",
            subtitle="Maha Shivaratri Brahmotsavam Series — February 2026",
            overview="A compilation of satsang transcripts from the Quantum Entanglement series.",
            chapters=[
                ChapterPlan(
                    chapter_number=i + 1,
                    day_number=t.day_number,
                    date=t.date,
                    chapter_title=f"Day {t.day_number}: {t.title[:60]}",
                )
                for i, t in enumerate(transcripts)
            ],
        )

    # Ensure chapter count matches transcript count
    if len(structure.chapters) != len(transcripts):
        structure.chapters = [
            ChapterPlan(
                chapter_number=i + 1,
                day_number=t.day_number,
                date=t.date,
                chapter_title=f"Day {t.day_number}: {t.title[:60]}",
            )
            for i, t in enumerate(transcripts)
        ]

    for ch in structure.chapters:
        safe_print(f"  Ch {ch.chapter_number}: {ch.chapter_title[:60]}")

    return structure


# ── Phase 3: Shastra Enrichment (parallel) ─────────────────────────────────

async def run_shastra_enrichment(
    transcripts: List[TranscriptFile],
    structure: BookStructure,
    tracer: Tracer,
    config: AppConfig,
    mcp_tools: Optional[List[Any]] = None,
) -> Dict[int, ShastraEnrichment]:
    """Find additional shastra pramanas for each chapter."""
    print("\n  Phase 3: Shastra Enrichment")
    print("  " + "-" * 50)

    enrichments: Dict[int, ShastraEnrichment] = {}

    async def enrich_one(ch_plan: ChapterPlan, transcript: TranscriptFile) -> ShastraEnrichment:
        # Extract a short content preview for context
        preview = transcript.content[:1500]

        agent = Agent(
            name=f"ShastraFinder (Ch.{ch_plan.chapter_number})",
            role=f"Find shastra pramanas for chapter {ch_plan.chapter_number}",
            model=build_agent_model("researcher", config),
            instructions=[
                f"Find 8-10 EXACT Sanskrit shastra pramanas for Day {ch_plan.day_number}.",
                f"Themes: {', '.join(ch_plan.key_themes[:5])}",
                "",
                "CRITICAL: Each pramana MUST include:",
                "  - sanskrit: the ACTUAL Sanskrit text in Devanagari or IAST transliteration",
                "    Example: 'ईशावास्यमिदं सर्वं यत्किञ्च जगत्यां जगत्'",
                "    Or IAST: 'īśāvāsyamidaṃ sarvaṃ yatkiñca jagatyāṃ jagat'",
                "    NOT just an English description. The actual verse text.",
                "  - source: exact reference (e.g. 'Isha Upanishad 1', 'Bhagavad Gita 2.47')",
                "  - translation: faithful English translation",
                "  - insert_after: a short phrase from the satsang content near where this verse fits",
                "",
                "Use MCP tools to search for verses. TOOL NAMES:",
                "  HYPHENS for Jnanalaya: search-books  resolve-book  get-chapters  read-chapter",
                "  UNDERSCORES for SPH: search_chapters  get_chapter  list_books",
                "  Call search_chapters(query='<theme>') to find relevant chapters with verses.",
                "",
                "Satsang content preview:",
                preview[:800],
                "",
                "Output a ShastraEnrichment with chapter_number and pramanas list.",
            ],
            tools=mcp_tools or [],
            markdown=False,
            tool_call_limit=3,
        )

        response = await tracer.traced_arun(
            agent,
            f"Find shastra pramanas for Day {ch_plan.day_number}: {ch_plan.chapter_title}",
            phase=f"Shastra-Ch{ch_plan.chapter_number}",
            output_schema=ShastraEnrichment,
        )

        if isinstance(response.content, ShastraEnrichment):
            result = response.content
        else:
            result = ShastraEnrichment(chapter_number=ch_plan.chapter_number)

        result.chapter_number = ch_plan.chapter_number
        safe_print(f"  Ch {ch_plan.chapter_number}: {len(result.pramanas)} pramanas found")
        return result

    tasks = [
        enrich_one(ch_plan, transcript)
        for ch_plan, transcript in zip(structure.chapters, transcripts)
    ]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    for result in results:
        if isinstance(result, Exception):
            safe_print(f"  WARNING: Enrichment error: {result}")
            continue
        enrichments[result.chapter_number] = result

    return enrichments


# ── Phase 3.5: Humor Stories + Extended Shastras (parallel) ────────────────

class HumorStory(BaseModel):
    """A short humor story in Living Enlightenment style."""
    chapter_number: int
    story_text: str = Field(..., description="The complete humor story in italics, 2-4 paragraphs")
    moral_bridge: str = Field("", description="1-2 sentence bridge from the joke to the chapter's teaching")


class ExtendedShastras(BaseModel):
    """Additional Sanskrit verses for a chapter."""
    chapter_number: int
    verses: List[dict] = Field(
        default_factory=list,
        description='List of {"sanskrit": ..., "source": ..., "translation": ..., "context": ...}',
    )


class AptaPramana(BaseModel):
    """Story from life of a Hindu saint/incarnation that fits a key truth."""
    chapter_number: int = 0
    stories: List[dict] = Field(
        default_factory=list,
        description='List of {"saint": ..., "story": ..., "truth_it_illustrates": ...}',
    )


class ScientificReference(BaseModel):
    """Scientific research study validating a key concept."""
    chapter_number: int = 0
    references: List[dict] = Field(
        default_factory=list,
        description='List of {"study": ..., "institution": ..., "finding": ..., "concept_validated": ...}',
    )


class BoxMessages(BaseModel):
    """Punchlines / powerful cognitions to be displayed as box messages."""
    chapter_number: int = 0
    messages: List[str] = Field(default_factory=list)


async def run_humor_and_shastras(
    transcripts: List[TranscriptFile],
    structure: BookStructure,
    tracer: Tracer,
    config: AppConfig,
    mcp_tools: Optional[List[Any]] = None,
) -> tuple:
    """Generate humor stories and find more Sanskrit verses for each chapter."""
    print("\n  Phase 3.5: Humor Stories + Extended Shastras")
    print("  " + "-" * 50)

    humor_results: Dict[int, HumorStory] = {}
    shastra_results: Dict[int, ExtendedShastras] = {}

    async def generate_humor(ch_plan: ChapterPlan, transcript: TranscriptFile) -> HumorStory:
        """Generate a Living Enlightenment-style humor story for one chapter."""
        preview = transcript.content[:600]

        agent = Agent(
            name=f"Humor (Ch.{ch_plan.chapter_number})",
            role=f"Write a humor story for chapter {ch_plan.chapter_number}",
            model=build_agent_model("writer", config),
            instructions=[
                "Write ONE short humorous story in the EXACT style of Living Enlightenment book.",
                "",
                "REFERENCE EXAMPLES from Living Enlightenment:",
                "",
                "Example 1:",
                "> *A disciple went to his master and said, 'Master, I am not able to meditate.",
                "> My legs ache. I feel distracted.' The master just said, 'It will pass.'*",
                "> *After two weeks, the disciple went back saying, 'I am able to meditate",
                "> beautifully. I feel so aware and blissful.' The master replied, 'It will pass.'*",
                "",
                "Example 2:",
                "> *Four men were in a private airplane. The tail caught fire. Only three",
                "> parachutes. The first grabbed one: 'My wife is waiting!' The second declared,",
                "> 'I am the most intelligent person on earth!' and jumped. An old man told the",
                "> last passenger, 'Take the last parachute.' The young man said, 'We can both go.",
                "> The most intelligent person just jumped with my backpack!'*",
                "",
                "Example 3:",
                "> *A boy told his teacher, 'Yesterday I killed three female and two male flies.'",
                "> The teacher asked, 'How did you know?' The boy replied, 'Three were on the",
                "> mirror and two were on the cigarette box!'*",
                "",
                "RULES:",
                "  - Story must be 2-4 paragraphs, written in ITALICS (*text*)",
                "  - Start with: 'A small story:' or 'Let me tell you a small story:'",
                "  - The story must be a JOKE or PARABLE with a punchline or twist",
                "  - It must relate to the chapter's theme (provided below)",
                "  - After the story, write 1-2 sentences bridging the humor to the teaching",
                "  - Keep it WARM, LIGHT, and RELATABLE — not preachy",
                "  - DO NOT write a long narrative — this is a SHORT joke/parable",
                "",
                f"Chapter theme: {ch_plan.chapter_title}",
                f"Key topics: {', '.join(ch_plan.key_themes[:3])}",
                f"Satsang preview: {preview[:300]}",
                "",
                "Output a HumorStory with story_text and moral_bridge.",
            ],
            tools=[],
            markdown=False,
        )

        response = await tracer.traced_arun(
            agent,
            f"Write a humor story for chapter {ch_plan.chapter_number}: {ch_plan.chapter_title}",
            phase=f"Humor-Ch{ch_plan.chapter_number}",
            output_schema=HumorStory,
        )

        if isinstance(response.content, HumorStory):
            result = response.content
        else:
            raw = str(response.content or "").strip()
            result = HumorStory(
                chapter_number=ch_plan.chapter_number,
                story_text=raw,
                moral_bridge="",
            )
        result.chapter_number = ch_plan.chapter_number
        safe_print(f"  Ch {ch_plan.chapter_number}: Humor story ({len(result.story_text.split())}w)")
        return result

    async def find_more_shastras(ch_plan: ChapterPlan, transcript: TranscriptFile) -> ExtendedShastras:
        """Find additional Sanskrit verses from MCP sources."""
        preview = transcript.content[:800]

        agent = Agent(
            name=f"ShastraExt (Ch.{ch_plan.chapter_number})",
            role=f"Find Sanskrit verses for chapter {ch_plan.chapter_number}",
            model=build_agent_model("researcher", config),
            instructions=[
                f"Find 8-10 EXACT Sanskrit shastra pramanas for chapter: {ch_plan.chapter_title}",
                f"Themes: {', '.join(ch_plan.key_themes[:4])}",
                "",
                "Each verse MUST have:",
                "  - sanskrit: ACTUAL Sanskrit in Devanagari or IAST",
                "    e.g. 'ईशावास्यमिदं सर्वं यत्किञ्च जगत्यां जगत्' or",
                "    'īśāvāsyamidaṃ sarvaṃ yatkiñca jagatyāṃ jagat'",
                "  - source: exact reference (e.g. 'Isha Upanishad 1', 'Bhagavad Gita 2.47')",
                "  - translation: English translation",
                "  - context: which topic in the satsang this verse supports",
                "",
                "Search in these sources:",
                "  - Bhagavad Gita, Yoga Sutras, Upanishads (Isha, Mandukya, Katha, Chandogya)",
                "  - Shiva Sutras, Vijnana Bhairava Tantra, Brahma Sutras",
                "  - Any Agama text relevant to the theme",
                "",
                "TOOL NAMES: search_chapters, get_chapter (SPH/underscores)",
                "  search-books, resolve-book, read-chapter (Jnanalaya/hyphens)",
                "",
                "Satsang preview:",
                preview[:500],
            ],
            tools=mcp_tools or [],
            markdown=False,
            tool_call_limit=3,
        )

        response = await tracer.traced_arun(
            agent,
            f"Find Sanskrit verses for chapter {ch_plan.chapter_number}",
            phase=f"ShastraExt-Ch{ch_plan.chapter_number}",
            output_schema=ExtendedShastras,
        )

        if isinstance(response.content, ExtendedShastras):
            result = response.content
        else:
            result = ExtendedShastras(chapter_number=ch_plan.chapter_number)
        result.chapter_number = ch_plan.chapter_number
        safe_print(f"  Ch {ch_plan.chapter_number}: {len(result.verses)} additional verses")
        return result

    # ── Apta Pramanas (stories of saints) ─────────────────────────────
    async def find_apta_pramanas(ch_plan: ChapterPlan, transcript: TranscriptFile) -> AptaPramana:
        preview = transcript.content[:500]
        agent = Agent(
            name=f"AptaPramana (Ch.{ch_plan.chapter_number})",
            role=f"Find saint stories for chapter {ch_plan.chapter_number}",
            model=build_agent_model("writer", config),
            instructions=[
                f"Find 2-3 stories from the lives of Hindu Incarnations and Saints",
                f"that illustrate the key truths in this chapter: {ch_plan.chapter_title}",
                f"Themes: {', '.join(ch_plan.key_themes[:3])}",
                "",
                "For each story provide:",
                "  - saint: name (e.g. Adi Shankara, Ramakrishna, Ramana Maharshi, Patanjali, Vyasa)",
                "  - story: the incident in 3-5 sentences — vivid, specific, not generic",
                "  - truth_it_illustrates: which teaching from the satsang this story supports",
                "",
                "Use REAL historical incidents, not invented ones.",
                "Output an AptaPramana with chapter_number and stories list.",
            ],
            tools=[],
            markdown=False,
        )
        response = await tracer.traced_arun(
            agent, f"Find saint stories for chapter {ch_plan.chapter_number}",
            phase=f"Apta-Ch{ch_plan.chapter_number}", output_schema=AptaPramana,
        )
        if isinstance(response.content, AptaPramana):
            result = response.content
        else:
            result = AptaPramana(chapter_number=ch_plan.chapter_number)
        result.chapter_number = ch_plan.chapter_number
        safe_print(f"  Ch {ch_plan.chapter_number}: {len(result.stories)} apta pramanas")
        return result

    # ── Scientific References ──────────────────────────────────────────
    async def find_scientific_refs(ch_plan: ChapterPlan, transcript: TranscriptFile) -> ScientificReference:
        preview = transcript.content[:500]
        agent = Agent(
            name=f"SciRef (Ch.{ch_plan.chapter_number})",
            role=f"Find scientific studies for chapter {ch_plan.chapter_number}",
            model=build_agent_model("writer", config),
            instructions=[
                f"Find 2-3 scientific research studies from reputed institutions",
                f"that validate key concepts in this chapter: {ch_plan.chapter_title}",
                f"Themes: {', '.join(ch_plan.key_themes[:3])}",
                "",
                "For each reference provide:",
                "  - study: title or description of the research",
                "  - institution: university/lab (e.g. 'Harvard Medical School', 'MIT', 'Max Planck Institute')",
                "  - finding: what the study found in 1-2 sentences",
                "  - concept_validated: which teaching from the satsang this validates",
                "",
                "Focus on: quantum physics, consciousness studies, meditation research,",
                "neuroscience, epigenetics, observer effect, entanglement experiments.",
                "Use REAL published research, not invented studies.",
                "Output a ScientificReference with chapter_number and references list.",
            ],
            tools=[],
            markdown=False,
        )
        response = await tracer.traced_arun(
            agent, f"Find scientific references for chapter {ch_plan.chapter_number}",
            phase=f"SciRef-Ch{ch_plan.chapter_number}", output_schema=ScientificReference,
        )
        if isinstance(response.content, ScientificReference):
            result = response.content
        else:
            result = ScientificReference(chapter_number=ch_plan.chapter_number)
        result.chapter_number = ch_plan.chapter_number
        safe_print(f"  Ch {ch_plan.chapter_number}: {len(result.references)} scientific refs")
        return result

    # ── Box Messages (punchlines) ──────────────────────────────────────
    async def extract_box_messages(ch_plan: ChapterPlan, transcript: TranscriptFile) -> BoxMessages:
        preview = transcript.content[:2000]
        agent = Agent(
            name=f"BoxMsg (Ch.{ch_plan.chapter_number})",
            role=f"Extract punchlines for chapter {ch_plan.chapter_number}",
            model=build_agent_model("editor", config),
            instructions=[
                f"Extract 5-8 powerful punchline statements from this satsang transcript",
                f"that can be displayed as highlighted box messages in the book.",
                "",
                "These should be:",
                "  - Short (1-2 sentences max)",
                "  - Powerful cognitions or declarations",
                "  - Quotable standalone truths",
                "  - Example: 'You are not a body having consciousness. You are consciousness having a body.'",
                "",
                "Extract ONLY from the actual transcript — do not invent new statements.",
                "",
                "Transcript preview:",
                preview,
                "",
                "Output a BoxMessages with chapter_number and messages list.",
            ],
            tools=[],
            markdown=False,
        )
        response = await tracer.traced_arun(
            agent, f"Extract box messages for chapter {ch_plan.chapter_number}",
            phase=f"BoxMsg-Ch{ch_plan.chapter_number}", output_schema=BoxMessages,
        )
        if isinstance(response.content, BoxMessages):
            result = response.content
        else:
            result = BoxMessages(chapter_number=ch_plan.chapter_number)
        result.chapter_number = ch_plan.chapter_number
        safe_print(f"  Ch {ch_plan.chapter_number}: {len(result.messages)} box messages")
        return result

    # Run ALL tasks in parallel
    all_tasks = []
    for ch_plan, transcript in zip(structure.chapters, transcripts):
        all_tasks.append(("humor", ch_plan, generate_humor(ch_plan, transcript)))
        all_tasks.append(("shastra", ch_plan, find_more_shastras(ch_plan, transcript)))
        all_tasks.append(("apta", ch_plan, find_apta_pramanas(ch_plan, transcript)))
        all_tasks.append(("sciref", ch_plan, find_scientific_refs(ch_plan, transcript)))
        all_tasks.append(("boxmsg", ch_plan, extract_box_messages(ch_plan, transcript)))

    results = await asyncio.gather(
        *[task for _, _, task in all_tasks],
        return_exceptions=True,
    )

    apta_results: Dict[int, AptaPramana] = {}
    sciref_results: Dict[int, ScientificReference] = {}
    boxmsg_results: Dict[int, BoxMessages] = {}

    for i, result in enumerate(results):
        task_type = all_tasks[i][0]
        ch_plan = all_tasks[i][1]
        if isinstance(result, Exception):
            safe_print(f"  WARNING: {task_type} Ch{ch_plan.chapter_number} failed: {result}")
            continue
        if task_type == "humor":
            humor_results[result.chapter_number] = result
        elif task_type == "shastra":
            shastra_results[result.chapter_number] = result
        elif task_type == "apta":
            apta_results[result.chapter_number] = result
        elif task_type == "sciref":
            sciref_results[result.chapter_number] = result
        elif task_type == "boxmsg":
            boxmsg_results[result.chapter_number] = result

    return humor_results, shastra_results, apta_results, sciref_results, boxmsg_results


# ── Phase 4: Format (per chapter) ──────────────────────────────────────────

# ── Sanskrit term preservation list ────────────────────────────────────────
# Words that should keep specific casing when converting from ALL CAPS
_SACRED_TERMS = {
    "paramashiva", "paramadvaita", "paramashivoham", "paramporul", "nithyananda",
    "shivaratri", "mahashivaratri", "paramashivaratri", "brahmotsavam",
    "shaktipada", "shaktinipada", "shaktinipata", "jnana", "rasavada",
    "upanishad", "mandukya", "isha", "katha", "chandogya", "taittiriya",
    "brihadaranyaka", "mundaka", "shvetashvatara", "kaivalya",
    "bhagavad", "gita", "yoga", "sutra", "patanjali", "veda", "vedanta",
    "agama", "purana", "dharma", "artha", "kama", "moksha", "purushartha",
    "atman", "brahman", "maya", "avidya", "samsara", "karma", "nirvana",
    "turiya", "turiyatita", "sushupti", "jagrat", "swapna",
    "ahamkara", "chitta", "buddhi", "manas", "antahkarana",
    "kundalini", "chakra", "nadi", "prana", "pranayama", "sushumna",
    "muladhara", "swadhishthana", "manipuraka", "anahata", "vishuddhi",
    "ajna", "sahasrara", "shakti", "shiva", "sadashiva", "nataraja",
    "deeksha", "diksha", "satsang", "sangha", "ashram", "gurukul",
    "kailasa", "mahakailasa", "bidadi", "varanasi", "kashi", "tiruvannamalai",
    "bhakti", "shraddha", "tapas", "samadhi", "nidhidhyasana", "manana",
    "shravana", "vichara", "viveka", "vairagya", "ananda", "anubhuti",
    "om", "aum", "namah", "sivaya", "soham", "shivoham", "tattvamasi",
    "sanatana", "hindu", "vedic", "agamic",
    "krishna", "arjuna", "vyasa", "shankara", "ramanuja", "ramakrishna",
    "buddha", "mahavira", "ramana", "maharshi",
    "panchakritya", "ishwara", "tathata", "avyakta", "vyakta",
    "srishti", "sthiti", "samhara", "tirobhava", "anugraha",
    "rajya", "sabha", "chit", "sat", "sakshi",
    "murti", "murthy", "lingam", "linga", "yagna", "yaga", "homa", "puja",
    "suktam", "stotram", "mantra", "shloka", "sloka",
}

# Build a title-case lookup
_SACRED_TITLE = {}
for _t in _SACRED_TERMS:
    _SACRED_TITLE[_t] = _t[0].upper() + _t[1:]  # simple title case


def _smart_case(text: str) -> str:
    """
    Convert ALL CAPS text to proper sentence case while preserving:
    - Sanskrit terms (title-cased from the sacred terms list)
    - Sanskrit verses with diacritics (kept as-is)
    - Proper nouns
    """
    # If text is NOT mostly uppercase, return as-is
    alpha_chars = [c for c in text if c.isalpha()]
    if not alpha_chars:
        return text
    upper_ratio = sum(1 for c in alpha_chars if c.isupper()) / len(alpha_chars)
    if upper_ratio < 0.7:
        return text  # not ALL CAPS, leave it

    # Check if this is a Sanskrit verse line (has diacritics)
    if re.search(r'[ĀĪŪṚṜḶḸĒŌṂḤŚṢṆṬḌĀĪŪṚṜḶḸĒŌṂḤ]', text):
        # Keep Sanskrit verses as-is (they need their exact casing)
        return text

    # Convert to lowercase first
    result = text.lower()

    # Title-case sacred terms
    for term, titled in _SACRED_TITLE.items():
        # Word boundary replacement
        result = re.sub(
            r'\b' + re.escape(term) + r'\b',
            titled,
            result,
            flags=re.IGNORECASE,
        )

    # Capitalize first letter of each sentence
    result = re.sub(r'(^|[.!?]\s+)([a-z])', lambda m: m.group(1) + m.group(2).upper(), result)

    # Capitalize "I" as standalone word
    result = re.sub(r"\bi\b", "I", result)

    # Capitalize first character of the whole string
    if result and result[0].islower():
        result = result[0].upper() + result[1:]

    return result


def _is_metadata_line(line: str) -> bool:
    """Check if a line is metadata (table header, empty table cell, etc.)."""
    stripped = line.strip()
    # Markdown table lines
    if stripped.startswith('|') and '|' in stripped[1:]:
        return True
    # Table separator
    if re.match(r'^[\s|:\-]+$', stripped):
        return True
    return False


def _is_sanskrit_verse(text: str) -> bool:
    """Detect if text is a Sanskrit verse/mantra."""
    # Has diacritics
    if re.search(r'[ĀĪŪṚṜḶḸĒŌṂḤŚṢṆṬḌāīūṛṝḷḹēōṃḥśṣṇṭḍ]', text):
        return True
    # Common verse patterns
    if re.match(r'^\s*(OM|AUM)\b', text, re.IGNORECASE):
        words = text.split()
        # If it's a short Om line or a longer mantra
        if len(words) <= 3 or any(w for w in words if re.search(r'[ĀĪŪṚṃḥśṣ]', w)):
            return True
    return False


def _detect_section_topic(para: str) -> str:
    """Try to detect a section topic from Swamiji's discourse markers."""
    lower = para.lower()
    patterns = [
        (r"paramashiva'?s direct message", "Paramashiva's Direct Message"),
        (r"paramasatya[s]?\s+about\s+(.+?)[\.\,]", None),  # dynamic
        (r"now listen.*truth", "The Truth Revealed"),
        (r"quantum entanglement", "Quantum Entanglement"),
        (r"quantum enlightenment", "Quantum Enlightenment"),
        (r"meditation\s+technique", "Meditation Technique"),
        (r"spiritual alchemy", "Spiritual Alchemy Process"),
        (r"let'?s start", ""),
        (r"blessings to all", ""),
    ]
    for pattern, label in patterns:
        match = re.search(pattern, lower)
        if match:
            if label is None and match.lastindex:
                return match.group(1).strip().title()
            return label or ""
    return ""


def _format_transcript_deterministic(
    content: str,
    ch_num: int,
    ch_title: str,
    date: str,
    occasion: str,
    enrichment: ShastraEnrichment,
    humor: Optional[HumorStory] = None,
    extended_shastras: Optional[ExtendedShastras] = None,
    apta: Optional[AptaPramana] = None,
    sciref: Optional[ScientificReference] = None,
    boxmsgs: Optional[BoxMessages] = None,
) -> str:
    """
    Deterministic formatter — NO LLM. Preserves every word.

    Converts ALL CAPS to sentence case, strips metadata tables,
    detects Sanskrit verses, adds section breaks, inserts shastra pramanas.
    """
    lines_out = []

    # Chapter heading
    lines_out.append(f"## Chapter {ch_num}: {ch_title}")
    lines_out.append("")
    lines_out.append(f"*{date} | {occasion}*")
    lines_out.append("")
    lines_out.append("---")
    lines_out.append("")

    # ── Step 1: Strip metadata tables ──────────────────────────────────
    raw_lines = content.split("\n")
    body_lines = []
    in_table = False
    for line in raw_lines:
        if _is_metadata_line(line):
            in_table = True
            continue
        if in_table and line.strip() == "":
            in_table = False
            continue
        in_table = False
        body_lines.append(line)

    text = "\n".join(body_lines)

    # ── Step 2: Split into paragraphs ──────────────────────────────────
    # The transcripts use "* " bullet points as paragraph markers
    # Split on bullet points
    paragraphs = []
    current = []

    for line in text.split("\n"):
        stripped = line.strip()

        # Bullet point = new paragraph
        if stripped.startswith("* ") or stripped.startswith("- "):
            if current:
                paragraphs.append(" ".join(current))
                current = []
            # Remove the bullet marker
            content_after_bullet = stripped[2:].strip()
            if content_after_bullet:
                current.append(content_after_bullet)
        elif stripped:
            current.append(stripped)
        else:
            if current:
                paragraphs.append(" ".join(current))
                current = []

    if current:
        paragraphs.append(" ".join(current))

    # ── Step 3: Process each paragraph ─────────────────────────────────
    shastra_inserted = set()
    para_count = 0
    section_count = 0
    last_was_section_break = True  # start after heading

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Detect if this is a Sanskrit verse
        if _is_sanskrit_verse(para):
            lines_out.append(f"> *{para}*")
            lines_out.append("")
            para_count += 1
            continue

        # Detect section transitions — add section heading
        upper_para = para.upper()
        is_discourse_marker = bool(re.match(
            r"^(LISTEN\.?|NOW,?\s*LISTEN|UNDERSTAND\.?|OM[\.\s]*OM|BLESSINGS|"
            r"PARAMASHIVA'?S DIRECT MESSAGE|PARAMASATYA)",
            upper_para,
        ))

        # Add section break for major transitions (not every "Listen")
        topic = _detect_section_topic(para)
        if topic and not last_was_section_break and para_count > 3:
            lines_out.append("")
            lines_out.append("---")
            lines_out.append("")
            lines_out.append(f"### {topic}")
            lines_out.append("")
            section_count += 1
            last_was_section_break = True
        elif is_discourse_marker and not last_was_section_break and para_count > 5:
            lines_out.append("")
            lines_out.append("---")
            lines_out.append("")
            last_was_section_break = True

        # Convert case and format as bullet point
        formatted_para = _smart_case(para)
        # Present as bullet point (editor's request: Bhagavan's utterances as bullets)
        if not formatted_para.startswith(('#', '>', '---', '![', '*')):
            lines_out.append(f"- {formatted_para}")
        else:
            lines_out.append(formatted_para)
        lines_out.append("")
        para_count += 1
        last_was_section_break = False

        # Insert box message every ~15 paragraphs
        if boxmsgs and boxmsgs.messages and para_count % 15 == 0:
            msg_idx = (para_count // 15 - 1) % len(boxmsgs.messages)
            if msg_idx < len(boxmsgs.messages):
                msg = boxmsgs.messages[msg_idx]
                lines_out.append(f"> **{msg}**")
                lines_out.append("")

        # Insert shastra pramanas near relevant content
        if enrichment.pramanas:
            para_lower = para.lower()
            for p in enrichment.pramanas:
                insert_key = p.get('insert_after', '')
                if insert_key and id(p) not in shastra_inserted:
                    if insert_key.lower() in para_lower:
                        sanskrit = p.get('sanskrit', '')
                        source = p.get('source', '')
                        translation = p.get('translation', '')
                        if sanskrit:
                            lines_out.append(f"> *{sanskrit}* — {source} — '{translation}'")
                            lines_out.append("")
                            shastra_inserted.add(id(p))

    # ── Insert humor story (after ~40% of the content) ──────────────
    if humor and humor.story_text:
        insert_pos = int(len(lines_out) * 0.4)
        # Find nearest paragraph break
        while insert_pos < len(lines_out) and lines_out[insert_pos].strip():
            insert_pos += 1
        humor_block = [
            "",
            "---",
            "",
            humor.story_text,
            "",
        ]
        if humor.moral_bridge:
            humor_block.append(humor.moral_bridge)
            humor_block.append("")
        humor_block.append("---")
        humor_block.append("")
        for i, line in enumerate(humor_block):
            lines_out.insert(insert_pos + i, line)

    # ── Append remaining shastra pramanas at the end ───────────────
    remaining = [p for p in (enrichment.pramanas or []) if id(p) not in shastra_inserted]
    all_extra_shastras = list(remaining)
    if extended_shastras and extended_shastras.verses:
        all_extra_shastras.extend(extended_shastras.verses)

    if all_extra_shastras:
        lines_out.append("")
        lines_out.append("---")
        lines_out.append("")
        lines_out.append("### Shastra Pramanas")
        lines_out.append("")
        for p in all_extra_shastras:
            sanskrit = p.get('sanskrit', '')
            source = p.get('source', '')
            translation = p.get('translation', '')
            context = p.get('context', '')
            if sanskrit:
                lines_out.append(f"> *{sanskrit}*")
                lines_out.append(f"> — {source}")
                if translation:
                    lines_out.append(f"> '{translation}'")
                if context:
                    lines_out.append(f"> (Context: {context})")
                lines_out.append("")

    # ── Apta Pramanas (stories of saints) ─────────────────────────────
    if apta and apta.stories:
        lines_out.append("")
        lines_out.append("---")
        lines_out.append("")
        lines_out.append("### Stories from the Lives of Saints (Apta Pramanas)")
        lines_out.append("")
        for s in apta.stories:
            saint = s.get('saint', '')
            story = s.get('story', '')
            truth = s.get('truth_it_illustrates', '')
            if story:
                lines_out.append(f"**{saint}**")
                lines_out.append("")
                lines_out.append(f"*{story}*")
                if truth:
                    lines_out.append(f"\n(Illustrates: {truth})")
                lines_out.append("")

    # ── Scientific References ──────────────────────────────────────────
    if sciref and sciref.references:
        lines_out.append("")
        lines_out.append("---")
        lines_out.append("")
        lines_out.append("### Scientific Research References")
        lines_out.append("")
        for r in sciref.references:
            study = r.get('study', '')
            institution = r.get('institution', '')
            finding = r.get('finding', '')
            concept = r.get('concept_validated', '')
            if study:
                lines_out.append(f"**{study}** ({institution})")
                lines_out.append(f"- Finding: {finding}")
                if concept:
                    lines_out.append(f"- Validates: {concept}")
                lines_out.append("")

    result = "\n".join(lines_out)

    # ── Post-processing ────────────────────────────────────────────────
    # Remove stray backslashes (from markdown escaping in source)
    result = result.replace("\\-", "-").replace("\\*", "*")
    result = re.sub(r'\\([^\\n])', r'\1', result)  # remove single backslash before any char except n

    # Convert square brackets to round braces for references
    result = re.sub(r'\[([^\]]*)\](?!\()', r'(\1)', result)  # [text] -> (text) but not [text](url)

    return result


async def run_formatting(
    transcripts: List[TranscriptFile],
    structure: BookStructure,
    enrichments: Dict[int, ShastraEnrichment],
    tracer: Tracer,
    config: AppConfig,
    humor_stories: Optional[Dict[int, HumorStory]] = None,
    extended_shastras: Optional[Dict[int, ExtendedShastras]] = None,
    apta_pramanas: Optional[Dict[int, AptaPramana]] = None,
    sci_references: Optional[Dict[int, ScientificReference]] = None,
    box_messages: Optional[Dict[int, BoxMessages]] = None,
) -> List[FormattedChapter]:
    """
    Format each transcript DETERMINISTICALLY — no LLM, no word loss.
    Every word from the original transcript is preserved.
    Enrichments (humor, shastras, apta, science, box messages) are INSERTED.
    """
    print("\n  Phase 4: Format (deterministic — zero word loss)")
    print("  " + "-" * 50)

    formatted: List[FormattedChapter] = []

    for ch_plan, transcript in zip(structure.chapters, transcripts):
        ch_num = ch_plan.chapter_number
        enrichment = enrichments.get(ch_num, ShastraEnrichment(chapter_number=ch_num))

        humor = (humor_stories or {}).get(ch_num)
        ext_shastras = (extended_shastras or {}).get(ch_num)
        apta = (apta_pramanas or {}).get(ch_num)
        sciref = (sci_references or {}).get(ch_num)
        boxmsgs = (box_messages or {}).get(ch_num)

        result = _format_transcript_deterministic(
            content=transcript.content,
            ch_num=ch_num,
            ch_title=ch_plan.chapter_title,
            date=ch_plan.date,
            occasion=ch_plan.occasion,
            enrichment=enrichment,
            humor=humor,
            extended_shastras=ext_shastras,
            apta=apta,
            sciref=sciref,
            boxmsgs=boxmsgs,
        )

        formatted_ch = FormattedChapter(
            chapter_number=ch_num,
            title=ch_plan.chapter_title,
            date=ch_plan.date,
            content_markdown=result,
            word_count=count_words(result),
            original_word_count=transcript.word_count,
        )
        formatted.append(formatted_ch)

        diff = formatted_ch.word_count - transcript.word_count
        sign = "+" if diff >= 0 else ""
        safe_print(
            f"  Ch {ch_num}: {transcript.word_count}w -> {formatted_ch.word_count}w ({sign}{diff}w)"
        )

    return formatted


# ── Phase 4.5: Grammar Correction (parallel) ──────────────────────────────

async def run_grammar_correction(
    formatted: List[FormattedChapter],
    tracer: Tracer,
    config: AppConfig,
) -> List[FormattedChapter]:
    """Fix grammatical errors in each chapter without altering meaning or content."""
    print("\n  Phase 4.5: Grammar Correction")
    print("  " + "-" * 50)

    async def correct_one(chapter: FormattedChapter) -> FormattedChapter:
        """
        Split by paragraph blocks (~1500 words each), correct each block.
        STRICT rule: if a corrected block is >3% shorter, reject it and keep original.
        """
        content = chapter.content_markdown

        # Split into paragraph blocks at double newlines
        paragraphs = content.split("\n\n")

        # Group paragraphs into chunks of ~1500 words
        chunks = []
        current_chunk = []
        current_wc = 0
        for para in paragraphs:
            para_wc = len(para.split())
            if current_wc + para_wc > 1500 and current_chunk:
                chunks.append("\n\n".join(current_chunk))
                current_chunk = [para]
                current_wc = para_wc
            else:
                current_chunk.append(para)
                current_wc += para_wc
        if current_chunk:
            chunks.append("\n\n".join(current_chunk))

        corrected_chunks = []
        for chunk in chunks:
            orig_wc = count_words(chunk)
            if orig_wc < 10:
                corrected_chunks.append(chunk)
                continue

            agent = Agent(
                name=f"Grammar (Ch.{chapter.chapter_number})",
                role="Fix grammar in satsang text",
                model=build_agent_model("editor", config),
                instructions=[
                    "EXTREMELY MINIMAL grammar polishing of satsang transcript.",
                    "Bhagavan's utterances are Ultimate sacred utterances.",
                    "",
                    "ONLY fix these:",
                    "  - Singular/plural mismatch (e.g. 'This 17 days' -> 'These 17 days')",
                    "  - Very glaring grammatical errors only",
                    "  - Add/remove punctuation for best readability",
                    "",
                    "DO NOT:",
                    "  - Replace words with 'better' words",
                    "  - Summarize, paraphrase, or condense anything",
                    "  - Remove, add, or reorder ANY content",
                    "  - Change Swamiji's speaking style — fragments, repetitions are INTENTIONAL",
                    "  - Alter Sanskrit terms, mantras, or markdown formatting (##, ---, > *, -)",
                    "",
                    "Return ONLY the corrected text, nothing else.",
                ],
                markdown=False,
            )

            response = await tracer.traced_arun(
                agent, chunk,
                phase=f"Grammar-Ch{chapter.chapter_number}",
            )

            corrected = str(response.content or "").strip()
            corrected_wc = count_words(corrected)

            # STRICT: reject if corrected lost more than 3% of words
            if corrected and corrected_wc >= orig_wc * 0.97:
                corrected_chunks.append(corrected)
            else:
                corrected_chunks.append(chunk)  # keep original

        corrected_content = "\n\n".join(corrected_chunks)
        new_wc = count_words(corrected_content)
        diff = new_wc - chapter.word_count
        sign = "+" if diff >= 0 else ""
        safe_print(f"  Ch {chapter.chapter_number}: {chapter.word_count}w -> {new_wc}w ({sign}{diff}w)")

        return FormattedChapter(
            chapter_number=chapter.chapter_number,
            title=chapter.title,
            date=chapter.date,
            content_markdown=corrected_content,
            word_count=new_wc,
            original_word_count=chapter.original_word_count,
        )

    # Run grammar correction in parallel
    tasks = [correct_one(ch) for ch in formatted]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    corrected = []
    for i, result in enumerate(results):
        if isinstance(result, Exception):
            safe_print(f"  WARNING: Grammar correction failed for Ch {formatted[i].chapter_number}: {result}")
            corrected.append(formatted[i])  # keep original
        else:
            corrected.append(result)

    return corrected


# ── Phase 5: QA ───────────────────────────────────────────────────────────

async def run_compilation_qa(
    transcripts: List[TranscriptFile],
    formatted: List[FormattedChapter],
    tracer: Tracer,
    config: AppConfig,
) -> CompilationQAResult:
    """Verify no content was lost and shastra pramanas are accurate."""
    print("\n  Phase 5: QA Review")
    print("  " + "-" * 50)

    # Build comparison data
    comparison_lines = []
    total_orig = 0
    total_final = 0
    for transcript, chapter in zip(transcripts, formatted):
        total_orig += transcript.word_count
        total_final += chapter.word_count
        # Check if key phrases from original appear in formatted
        # Take 10 random phrases from the original
        orig_words = transcript.content.split()
        sample_phrases = []
        for start in range(0, len(orig_words), max(1, len(orig_words) // 10)):
            phrase = " ".join(orig_words[start:start + 6])
            sample_phrases.append(phrase)

        missing = [p for p in sample_phrases[:10] if p.lower() not in chapter.content_markdown.lower()]

        comparison_lines.append(
            f"Ch {chapter.chapter_number} (Day {transcript.day_number}):\n"
            f"  Original: {transcript.word_count}w | Formatted: {chapter.word_count}w\n"
            f"  Sample phrases missing: {len(missing)}/{len(sample_phrases[:10])}\n"
            f"  Missing samples: {missing[:3] if missing else 'none'}\n"
        )

    agent = Agent(
        name="Compilation QA",
        role="Verify the compiled book preserves all original content.",
        model=build_agent_model("qa", config),
        instructions=[
            "You are the QA reviewer for a satsang transcript compilation.",
            "",
            "CRITICAL CHECK: Every word from the original transcript MUST be in the formatted chapter.",
            "The formatted chapter may have ADDITIONAL content (shastra pramanas, section breaks)",
            "but must NOT have LESS content than the original.",
            "",
            "CHECKS:",
            "  1. Word count: formatted should be ~= original (metadata removal causes -3 to -5% which is OK)",
            "     Formatted may be slightly lower due to stripped metadata tables — this is ACCEPTABLE.",
            "  2. Key phrases preserved: core spoken content must be present",
            "     (metadata lines like 'SATSANG TITLE | ...' are intentionally removed — ignore those)",
            "  3. Shastra pramanas: added verses must be in Sanskrit (Devanagari or IAST) with source + translation",
            "  4. Structure: chapters should have proper headings (##, ###) and section breaks (---)",
            "  5. Grammar: check for grammatical errors. Flag but do not fail for Swamiji's intentional",
            "     speech patterns ('Listen.', 'Understand.', fragments). Only flag actual errors.",
            "  6. No stray backslashes (\\) in the output",
            "  7. References use round braces () not square brackets []",
            "",
            "COMPARISON DATA:",
            "\n".join(comparison_lines),
            "",
            f"Total original words: {total_orig}",
            f"Total formatted words: {total_final}",
            f"Difference: {total_final - total_orig} words ({'+' if total_final >= total_orig else ''}{((total_final - total_orig) / total_orig * 100):.1f}%)",
            "",
            "APPROVE if formatted word count >= original for every chapter and no content is missing.",
            "REJECT if any chapter lost content.",
        ],
        markdown=False,
    )

    response = await tracer.traced_arun(
        agent,
        "Verify the compilation preserves all original transcript content.",
        phase="Compilation-QA",
        output_schema=CompilationQAResult,
    )

    if isinstance(response.content, CompilationQAResult):
        qa_result = response.content
    else:
        # Auto-check based on word counts
        issues = []
        for transcript, chapter in zip(transcripts, formatted):
            if chapter.word_count < transcript.word_count * 0.9:
                issues.append(
                    f"Ch {chapter.chapter_number}: lost content "
                    f"({transcript.word_count}w -> {chapter.word_count}w)"
                )
        qa_result = CompilationQAResult(
            approved=len(issues) == 0,
            chapters_checked=len(formatted),
            words_original=total_orig,
            words_final=total_final,
            issues=issues,
        )

    safe_print(f"  Original: {qa_result.words_original:,}w")
    safe_print(f"  Final:    {qa_result.words_final:,}w")
    safe_print(f"  Status:   {'APPROVED' if qa_result.approved else 'ISSUES FOUND'}")
    if qa_result.issues:
        for issue in qa_result.issues:
            safe_print(f"    - {issue}")

    return qa_result


# ── Phase 6: Design (.docx) ───────────────────────────────────────────────

def build_compilation_docx(
    structure: BookStructure,
    chapters: List[FormattedChapter],
    output_path: Path,
) -> None:
    """Build styled .docx from formatted chapters."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("  WARNING: python-docx not installed. Skipping .docx generation.")
        return

    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Palatino Linotype"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)

    for level, (name, size, bold, color) in {
        0: ("Title", 28, True, RGBColor(0x1A, 0x1A, 0x2E)),
        1: ("Heading 1", 22, True, RGBColor(0x1A, 0x1A, 0x2E)),
        2: ("Heading 2", 14, True, RGBColor(0x0F, 0x34, 0x60)),
    }.items():
        try:
            h = doc.styles[name]
        except KeyError:
            h = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        h.font.name = "Georgia"
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.color.rgb = color

    # Title page
    tp = doc.add_paragraph()
    tp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = tp.add_run(structure.book_title)
    run.font.name = "Georgia"
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    if structure.subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sr = sp.add_run(structure.subtitle)
        sr.font.name = "Georgia"
        sr.font.size = Pt(16)
        sr.italic = True

    ap = doc.add_paragraph()
    ap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ar = ap.add_run("\n\nThe SPH Bhagwan Sri Nithyananda Paramashivam")
    ar.font.name = "Georgia"
    ar.font.size = Pt(14)

    doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents", level=1)
    for ch in chapters:
        doc.add_paragraph(f"Chapter {ch.chapter_number}: {ch.title}", style="Normal")
    doc.add_page_break()

    # Chapters
    for ch in chapters:
        doc.add_heading(f"Chapter {ch.chapter_number}: {ch.title}", level=1)
        # Parse markdown into paragraphs
        for line in ch.content_markdown.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("## "):
                pass  # skip chapter heading (already added)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=2)
            elif stripped.startswith("> "):
                try:
                    doc.add_paragraph(stripped[2:], style="Quote")
                except KeyError:
                    p = doc.add_paragraph(stripped[2:])
                    p.runs[0].italic = True if p.runs else None
            elif stripped.startswith("---"):
                doc.add_paragraph("―" * 30)
            elif stripped.startswith("*") and stripped.endswith("*"):
                p = doc.add_paragraph()
                r = p.add_run(stripped.strip("*"))
                r.italic = True
            else:
                doc.add_paragraph(stripped)
        doc.add_page_break()

    doc.save(str(output_path))


# ── Main Pipeline ──────────────────────────────────────────────────────────

async def compile_satsang_book(
    transcript_dir: str,
    config_path: Optional[str] = None,
) -> Path:
    """Run the full satsang compilation pipeline."""
    app_config = reload_config(config_path)
    transcript_path = Path(transcript_dir)

    if not transcript_path.exists():
        raise FileNotFoundError(f"Transcript directory not found: {transcript_path}")

    # Phase 1: Ingest
    transcripts = ingest_transcripts(transcript_path)

    # Setup output
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = ensure_dir(Path(app_config.workflow.output_dir) / f"satsang_compilation_{timestamp}")
    tracer = Tracer(output_dir=output_dir, verbose=True)

    total_words = sum(t.word_count for t in transcripts)
    print(f"\n{'=' * 72}")
    print(f"  Satsang Compilation Pipeline")
    print(f"  Transcripts: {len(transcripts)}")
    print(f"  Total words: {total_words:,}")
    print(f"  Output: {output_dir}")
    print(f"{'=' * 72}")

    # Save raw transcripts manifest
    write_json(output_dir / "manifest.json", [t.model_dump() for t in transcripts])

    # Connect MCP for shastra enrichment
    mcp_tools: List[Any] = []
    mcp_connections: List[MCPTools] = []
    print("\n  Connecting to MCP sources for shastra enrichment...")
    for source in app_config.mcp_sources:
        try:
            mcp = MCPTools(transport=source.transport, url=source.url, timeout_seconds=120)
            await mcp.connect()
            mcp_tools.append(mcp)
            mcp_connections.append(mcp)
            safe_print(f"  Connected: {source.name}")
        except Exception as exc:
            safe_print(f"  WARNING: {source.name} failed: {exc}")

    try:
        # Phase 1.5: Content Cleanup (deterministic)
        cleaned_transcripts, cleanup_reports = run_content_cleanup(transcripts)
        write_json(output_dir / "cleanup_reports.json", [r.model_dump() for r in cleanup_reports])
        # Save cleaned transcripts
        clean_dir = ensure_dir(output_dir / "cleaned_transcripts")
        for ct in cleaned_transcripts:
            write_text(clean_dir / f"day_{ct.day_number:02d}_cleaned.md", ct.content)

        # Phase 2: Structure (uses cleaned transcripts)
        structure = await run_structurer(cleaned_transcripts, tracer, app_config)
        write_json(output_dir / "structure.json", structure.model_dump())

        # Phase 3: Shastra Enrichment
        enrichments = await run_shastra_enrichment(
            cleaned_transcripts, structure, tracer, app_config, mcp_tools
        )
        write_json(output_dir / "enrichments.json", {k: v.model_dump() for k, v in enrichments.items()})

        # Phase 3.5: Humor + Shastras + Apta Pramanas + Science + Box Messages
        humor_stories, extended_shastras, apta_pramanas, sci_references, box_messages = (
            await run_humor_and_shastras(
                cleaned_transcripts, structure, tracer, app_config, mcp_tools
            )
        )
        write_json(output_dir / "humor_stories.json", {k: v.model_dump() for k, v in humor_stories.items()})
        write_json(output_dir / "extended_shastras.json", {k: v.model_dump() for k, v in extended_shastras.items()})
        write_json(output_dir / "apta_pramanas.json", {k: v.model_dump() for k, v in apta_pramanas.items()})
        write_json(output_dir / "sci_references.json", {k: v.model_dump() for k, v in sci_references.items()})
        write_json(output_dir / "box_messages.json", {k: v.model_dump() for k, v in box_messages.items()})

        # Phase 4: Format (deterministic)
        formatted = await run_formatting(
            cleaned_transcripts, structure, enrichments, tracer, app_config,
            humor_stories=humor_stories, extended_shastras=extended_shastras,
            apta_pramanas=apta_pramanas, sci_references=sci_references,
            box_messages=box_messages,
        )
        fmt_dir = ensure_dir(output_dir / "chapters_raw")
        for ch in formatted:
            write_text(fmt_dir / f"ch_{ch.chapter_number:02d}_{slugify(ch.title)}.md", ch.content_markdown)

        # Phase 4.5: Grammar Correction
        corrected = await run_grammar_correction(formatted, tracer, app_config)
        corrected_dir = ensure_dir(output_dir / "chapters")
        for ch in corrected:
            write_text(corrected_dir / f"ch_{ch.chapter_number:02d}_{slugify(ch.title)}.md", ch.content_markdown)

        # Phase 5: QA
        qa_result = await run_compilation_qa(transcripts, corrected, tracer, app_config)
        write_json(output_dir / "qa_result.json", qa_result.model_dump())

        # Phase 6: Design
        print("\n  Phase 6: Design")
        print("  " + "-" * 50)
        docx_path = output_dir / f"{slugify(structure.book_title)}.docx"
        build_compilation_docx(structure, corrected, docx_path)
        safe_print(f"  DOCX saved: {docx_path}")

        # Combined markdown
        md_path = output_dir / f"{slugify(structure.book_title)}.md"
        md_parts = [f"# {structure.book_title}\n\n*{structure.subtitle}*\n\n---\n"]
        for ch in corrected:
            md_parts.append(ch.content_markdown)
            md_parts.append("\n\n---\n\n")
        write_text(md_path, "\n".join(md_parts))
        safe_print(f"  Markdown saved: {md_path}")

    finally:
        for mcp in mcp_connections:
            try:
                await mcp.close()
            except Exception:
                pass

    # Summary
    print(tracer.summary())

    final_words = sum(ch.word_count for ch in corrected)
    print(f"{'=' * 72}")
    print(f"  Compilation complete!")
    print(f"  Original: {total_words:,} words")
    print(f"  Final:    {final_words:,} words")
    print(f"  Chapters: {len(formatted)}")
    print(f"  QA:       {'APPROVED' if qa_result.approved else 'ISSUES'}")
    print(f"  Output:   {output_dir}")
    print(f"{'=' * 72}\n")

    return docx_path


# ── CLI ────────────────────────────────────────────────────────────────────

async def main():
    transcript_dir = "transcripts_feb2026"

    # Parse CLI args
    args = sys.argv[1:]
    i = 0
    while i < len(args):
        if args[i] == "--dir" and i + 1 < len(args):
            transcript_dir = args[i + 1]
            i += 2
        else:
            i += 1

    await compile_satsang_book(transcript_dir)


if __name__ == "__main__":
    asyncio.run(main())
